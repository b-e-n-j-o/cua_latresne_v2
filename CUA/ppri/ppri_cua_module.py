# -*- coding: utf-8 -*-
"""
ppri_cua_module.py
Module de génération de rapports PPRI pour documents CUA (Word, texte brut).
Fonctions dédiées à la production de contenu textuel et tableaux.
"""

import geopandas as gpd
from sqlalchemy import text
# --- Import du moteur SQL global défini dans ppri_analyse_tolerance
try:
    from CUA.ppri.ppri_analyse_tolerance import engine as ENGINE_PPRI
    if ENGINE_PPRI is None:
        print("⚠️ ENGINE_PPRI est None — tentative de reconstruction manuelle.")
        from dotenv import load_dotenv
        import os
        from sqlalchemy import create_engine
        load_dotenv()
        HOST = os.getenv("SUPABASE_HOST")
        DB = os.getenv("SUPABASE_DB")
        USER = os.getenv("SUPABASE_USER")
        PWD = os.getenv("SUPABASE_PASSWORD")
        PORT = os.getenv("SUPABASE_PORT", 5432)
        ENGINE_PPRI = create_engine(f"postgresql+psycopg2://{USER}:{PWD}@{HOST}:{PORT}/{DB}")
    else:
        print("✅ ENGINE_PPRI importé depuis ppri_analyse_tolerance.")
except Exception as e:
    print(f"⚠️ Impossible d'importer ENGINE_PPRI : {e}")
    ENGINE_PPRI = None


# --- Import du module altimétrique IGN pour le paragraphe NGF général
try:
    from CUA.ppri.cote_ngf import cote_ngf_parcelle
except ImportError:
    cote_ngf_parcelle = None
    print("⚠️ Module cote_ngf non disponible : paragraphe altimétrique IGN désactivé.")

# 🔌 Import de la connexion existante depuis ppri_analyse_tolerance
try:
    from CUA.ppri.ppri_analyse_tolerance import engine as ppri_engine
    engine = ppri_engine
    print("✅ Engine PPRI importé depuis ppri_analyse_tolerance.")
except Exception as e:
    engine = None
    print(f"⚠️ Impossible d'importer l'engine PPRI : {e}")


# ============================================================
# 📝 GÉNÉRATION DE RAPPORTS POUR CUA
# ============================================================
def generer_rapport_cua(resultats):
    """
    Génère un rapport PPRI au format texte brut prêt à insérer dans le CUA (Article PM1).
    
    Args:
        resultats (dict): Dictionnaire retourné par analyser_ppri_corrige() ou analyser_ppri_tolerance()
                          Doit contenir:
                          - parcelle: {surface_m2, ...}
                          - zone_dominante: {nom, pourcentage, surface_m2, texte}
                          - zones_avec_regles: [{nom, pourcentage, surface_m2, texte}, ...]
                          - cas_multizone: bool
    
    Returns:
        str: Texte brut formaté, prêt à insérer dans le document CUA
    """
    
    dominante = resultats["zone_dominante"]
    zones = resultats["zones_avec_regles"]

    lignes = []
    lignes.append("Le PPRI combine zonage réglementaire et cotes de seuil (altitudes NGF de référence intégrant le changement climatique à horizon 2100).")
    lignes.append("")
    lignes.append("Répartition surfacique :")
    lignes.append(f"La parcelle de {round(resultats['parcelle']['surface_m2'])} m² est répartie comme suit :")
    lignes.append("")

    for z in zones:
        lignes.append(f"- {z['nom']} : {z['pourcentage']} % ({z['surface_m2']} m²)")

    if resultats["cas_multizone"]:
        lignes.append("")
        lignes.append("⚠️ Cas Multi-zones :")
        lignes.append("L'unité foncière est concernée par plusieurs zones réglementaires. L'article A – sec. IV – § c du règlement du PPRI prévoit trois options d'application pour les projets :")
        lignes.append("1. Application stricte : adapter chaque partie du projet aux prescriptions de la zone qui la concerne.")
        lignes.append("2. Règle la plus restrictive : appliquer les prescriptions de la zone la plus contraignante à l'ensemble du projet.")
        lignes.append("3. Implantation : implanter le projet entièrement dans la zone la moins exposée, si possible.")
        lignes.append("")
        lignes.append("Pour les bâtiments existants à cheval sur plusieurs zones (hors grenat et rouge non urbanisée), il est possible d'appliquer la réglementation de la zone la moins contraignante si quatre conditions sont remplies :")
        lignes.append("- la zone la moins contraignante est majoritaire en surface ;")
        lignes.append("- l'accès aux parties en zone plus exposée se fait uniquement depuis la zone moins exposée ;")
        lignes.append("- pas d'augmentation de vulnérabilité (planchers au-dessus de la cote de seuil) ;")
        lignes.append("- la partie la plus exposée ne communique pas avec l'extérieur sous la cote de seuil.")
        lignes.append("")

    # 🔹 Cotes de seuil PPRI
    cotes = resultats.get("cotes_ngf", [])
    if cotes:
        min_cote = round(min(cotes), 2)
        max_cote = round(max(cotes), 2)
        if abs(max_cote - min_cote) < 0.05:
            lignes.append(f"Cote de seuil relevée : {min_cote} m NGF.")
        else:
            lignes.append(f"Cotes de seuil relevées : entre {min_cote} m et {max_cote} m NGF.")
        lignes.append("")

    # 🔹 Paragraphe altimétrique IGN
    if resultats.get("paragraphe_altitude"):
        lignes.append(resultats["paragraphe_altitude"])
        lignes.append("")

    lignes.append("Prescriptions détaillées :")
    lignes.append("Les textes suivants sont applicables aux différentes parties de la parcelle :")
    lignes.append("")

    for z in zones:
        lignes.append(f"{z['nom'].upper()} ({z['pourcentage']} %)")
        lignes.append("-" * 60)
        lignes.append(z["texte"].replace("\n", " ").strip())
        lignes.append("")

    lignes.append("Glossaire :")
    lignes.append("- Aléa : probabilité qu'une inondation se produise avec une certaine intensité.")
    lignes.append("- Cote de seuil : hauteur minimale du premier plancher habitable pour être hors d'eau.")
    lignes.append("- Changement de destination : modification de l'usage d'un bâtiment (ex. commerce → logement).")
    lignes.append("- Vulnérabilité : degré d'exposition des personnes et biens au risque.")
    lignes.append("")
    lignes.append("Références :")
    lignes.append("- Arrêté préfectoral du 23 février 2022 (PPRI de Latresne).")
    lignes.append("- Règlement complet : https://www.mairie-latresne.fr/risques-inondations/")
    lignes.append("")

    return "\n".join(lignes)


def generer_rapport_cua_structure(resultats):
    """
    Variante structurée retournant un dictionnaire pour plus de flexibilité.
    
    Args:
        resultats (dict): Résultats de l'analyse PPRI
    
    Returns:
        dict: {
            "introduction": str,
            "repartition": str,
            "multizone": str | None,
            "prescriptions": list[{zone: str, texte: str}],
            "glossaire": str,
            "references": str,
            "texte_complet": str
        }
    """
    
    zones = resultats["zones_avec_regles"]
    
    # Introduction
    introduction = "Le PPRI combine zonage réglementaire et cotes de seuil (altitudes NGF de référence intégrant le changement climatique à horizon 2100)."
    
    # Répartition
    repartition_lignes = [
        "Répartition surfacique :",
        f"La parcelle de {round(resultats['parcelle']['surface_m2'])} m² est répartie comme suit :",
        ""
    ]
    for z in zones:
        repartition_lignes.append(f"- {z['nom']} : {z['pourcentage']} % ({z['surface_m2']} m²)")
    
    repartition = "\n".join(repartition_lignes)
    
    # Multi-zones
    multizone = None
    if resultats["cas_multizone"]:
        multizone_lignes = [
            "⚠️ Cas Multi-zones :",
            "L'unité foncière est concernée par plusieurs zones réglementaires. L'article A – sec. IV – § c du règlement du PPRI prévoit trois options d'application pour les projets :",
            "1. Application stricte : adapter chaque partie du projet aux prescriptions de la zone qui la concerne.",
            "2. Règle la plus restrictive : appliquer les prescriptions de la zone la plus contraignante à l'ensemble du projet.",
            "3. Implantation : implanter le projet entièrement dans la zone la moins exposée, si possible.",
            "",
            "Pour les bâtiments existants à cheval sur plusieurs zones (hors grenat et rouge non urbanisée), il est possible d'appliquer la réglementation de la zone la moins contraignante si quatre conditions sont remplies :",
            "- la zone la moins contraignante est majoritaire en surface ;",
            "- l'accès aux parties en zone plus exposée se fait uniquement depuis la zone moins exposée ;",
            "- pas d'augmentation de vulnérabilité (planchers au-dessus de la cote de seuil) ;",
            "- la partie la plus exposée ne communique pas avec l'extérieur sous la cote de seuil."
        ]
        multizone = "\n".join(multizone_lignes)
    
    # Prescriptions
    prescriptions = []
    for z in zones:
        prescriptions.append({
            "zone": f"{z['nom'].upper()} ({z['pourcentage']} %)",
            "texte": z["texte"].replace("\n", " ").strip()
        })
    
    # Glossaire
    glossaire = "\n".join([
        "Glossaire :",
        "- Aléa : probabilité qu'une inondation se produise avec une certaine intensité.",
        "- Cote de seuil : hauteur minimale du premier plancher habitable pour être hors d'eau.",
        "- Changement de destination : modification de l'usage d'un bâtiment (ex. commerce → logement).",
        "- Vulnérabilité : degré d'exposition des personnes et biens au risque."
    ])
    
    # Références
    references = "\n".join([
        "Références :",
        "- Arrêté préfectoral du 23 février 2022 (PPRI de Latresne).",
        "- Règlement complet : https://www.mairie-latresne.fr/risques-inondations/"
    ])
    
    # Texte complet
    texte_complet = generer_rapport_cua(resultats)
    
    return {
        "introduction": introduction,
        "repartition": repartition,
        "multizone": multizone,
        "prescriptions": prescriptions,
        "glossaire": glossaire,
        "references": references,
        "texte_complet": texte_complet
    }


# ============================================================
# 🧩 WRAPPER COMPATIBILITÉ (ancien nom de fonction)
# ============================================================
def analyser_ppri_corrige(geom_wkt=None, section=None, numero=None, code_insee=None, engine=None, **kwargs):
    """
    Wrapper rétrocompatible pour cua_builder.py
    → exécute analyser_ppri_tolerance() et ajoute :
       - cotes de seuil PPRI (PostGIS)
       - paragraphe altimétrique IGN (API Altimétrie)
    
    Accepte désormais :
    - geom_wkt : géométrie WKT directe (prioritaire)
    - section + numero : identifiants parcellaires (fallback)
    """
    print("⚙️ [Compatibilité] Appel de analyser_ppri_corrige() redirigé vers analyser_ppri_tolerance()")
    from CUA.ppri.ppri_analyse_tolerance import analyser_ppri_tolerance, engine as ppri_engine
    
    # ✅ Utiliser ENGINE_PPRI si engine n'est pas fourni
    if engine is None:
        engine = ENGINE_PPRI or ppri_engine
    
    if engine is None:
        raise RuntimeError("❌ Moteur SQL introuvable pour l'analyse PPRI")

    # --- Exécution de l'analyse de base
    resultats = analyser_ppri_tolerance(
        geom_wkt=geom_wkt,
        section=section,
        numero=numero,
        code_insee=code_insee,
        ppri_table="latresne.pm1_detaillee_gironde",
        engine=engine  # ✅ passage explicite
    )

    if not resultats or "zones" not in resultats:
        return {}

    # Construction simplifiée pour compatibilité avec cua_builder
    cons = resultats["zones"]["ppri_conservees"]
    surface_parcelle = resultats["parcelle"]["surface"]
    
    # ✅ Grouper les fragments par codezone (fusionner géométriquement)
    zones_groupees = cons.dissolve(by='codezone', aggfunc={
        'reglementation': 'first'  # Prendre la première réglementation
    }).reset_index()
    
    zones_liste = []
    for _, row in zones_groupees.iterrows():
        surface_zone = row.geometry.area
        zones_liste.append({
            "nom": str(row.codezone),
            "pourcentage": round((surface_zone / surface_parcelle) * 100, 2),
            "surface_m2": round(surface_zone, 2),
            "texte": str(row.reglementation or "").strip()
        })

    surface_totale = sum(z["surface_m2"] for z in zones_liste)
    zone_dominante = max(zones_liste, key=lambda z: z["surface_m2"], default={})

    # ============================================================
    # 🧮 ÉTAPE 1 — COTES DE SEUIL PPRI (colonne codezone = valeur NGF)
    # ============================================================
    cotes_ngf = []
    try:
        print("🌊 Recherche des cotes de seuil PPRI (colonne codezone)...")
        sql_cote = """
            WITH parcelle AS (
                SELECT ST_GeomFromText(:wkt, 2154) AS geom
            )
            SELECT c.codezone AS cote_ngf
            FROM latresne.cote_de_seuil_ppri c, parcelle p
            WHERE ST_Intersects(c.geom_2154, p.geom);
        """
        with engine.connect() as conn:
            res = conn.execute(text(sql_cote), {"wkt": geom_wkt})
            cotes_ngf = [float(r[0]) for r in res.fetchall() if r[0] is not None]

        if cotes_ngf:
            print(f"✅ {len(cotes_ngf)} cotes de seuil trouvées : {cotes_ngf}")
        else:
            print("⚠️ Aucune cote de seuil trouvée pour cette unité foncière.")
    except Exception as e:
        print(f"⚠️ Erreur chargement cotes de seuil PPRI : {e}")

    # ============================================================
    # 🧭 ÉTAPE 2 — COTE ALTIMÉTRIQUE IGN (API Altimétrie)
    # ============================================================
    paragraphe_altitude = ""
    if geom_wkt and cote_ngf_parcelle:
        try:
            print("📡 Calcul altimétrique global de l'unité foncière (API IGN)...")
            paragraphe_altitude = cote_ngf_parcelle(geom_wkt)
            print("✅ Paragraphe altimétrique IGN intégré.")
        except Exception as e:
            print(f"⚠️ Impossible de générer le paragraphe altimétrique IGN : {e}")
    else:
        print("ℹ️ Calcul altimétrique IGN désactivé (géométrie WKT manquante).")

    # ============================================================
    # 🔧 STRUCTURE FINALISÉE POUR LE BUILDER
    # ============================================================
    resultats_cua = {
        "parcelle": {"surface_m2": int(round(surface_totale))},
        "zone_dominante": zone_dominante,
        "zones_avec_regles": zones_liste,
        "cas_multizone": len(zones_liste) > 1,
        "cotes_ngf": cotes_ngf,
        "paragraphe_altitude": paragraphe_altitude
    }

    print(f"✅ Zones distinctes : {len(zones_liste)}")
    print(f"✅ Zone dominante : {zone_dominante.get('nom')} ({zone_dominante.get('pourcentage', 0)}%)")
    return resultats_cua


# ============================================================
# 🧩 WRAPPER COMPATIBILITÉ POUR generer_rapport_cua_avec_table
# ============================================================
def generer_rapport_cua_avec_table(doc, resultats):
    """
    Wrapper de compatibilité pour cua_builder.py.
    Génère un tableau PPRI synthétique dans le document DOCX.
    Utilise la logique de generer_rapport_cua() pour le texte principal.
    """
    from docx.shared import Pt
    from docx import Document

    # ✅ Structure simplifiée compatible avec le CUA builder
    zones = resultats.get("zones_avec_regles", [])
    dominante = resultats.get("zone_dominante", {})
    surface_totale = int(round(resultats.get("parcelle", {}).get("surface_m2", 0)))

    # Titre et introduction
    doc.add_paragraph("Le PPRI combine zonage réglementaire et cotes de seuil (altitudes NGF de référence intégrant le changement climatique à horizon 2100).")
    doc.add_paragraph(f"La zone dominante est {dominante.get('nom', 'inconnue')} ({dominante.get('pourcentage', 0)} % de la parcelle).")
    doc.add_paragraph(f"La parcelle de {surface_totale} m² est répartie comme suit :")

    # Création du tableau
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Zone").bold = True
    hdr_cells[1].paragraphs[0].add_run("Pourcentage").bold = True
    hdr_cells[2].paragraphs[0].add_run("Surface (m²)").bold = True

    for z in zones:
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(z["nom"])
        row_cells[1].paragraphs[0].add_run(f"{z['pourcentage']} %")
        row_cells[2].paragraphs[0].add_run(str(int(round(z["surface_m2"], 0))))

    doc.add_paragraph("")  # espacement

    # 🔹 Cotes de seuil PPRI (si dispo)
    cotes = resultats.get("cotes_ngf", [])
    if cotes:
        min_cote = round(min(cotes), 2)
        max_cote = round(max(cotes), 2)
        if abs(max_cote - min_cote) < 0.05:
            doc.add_paragraph(f"Cote de seuil relevée : {min_cote} m NGF.")
        else:
            doc.add_paragraph(f"Cotes de seuil relevées : entre {min_cote} m et {max_cote} m NGF.")
        doc.add_paragraph("")

    # 🔹 Paragraphe altimétrique IGN (si dispo)
    if resultats.get("paragraphe_altitude"):
        doc.add_paragraph(resultats["paragraphe_altitude"])
        doc.add_paragraph("")

    # 🔹 Prescriptions réglementaires
    doc.add_paragraph("Prescriptions détaillées :", style="List Bullet")
    for z in zones:
        doc.add_paragraph(f"{z['nom'].upper()} ({z['pourcentage']} %) :", style="List Bullet")
        doc.add_paragraph(z["texte"].replace("\n", " ").strip())

    doc.add_paragraph("")
    doc.add_paragraph("Références :")
    doc.add_paragraph("- Arrêté préfectoral du 23 février 2022 (PPRI de Latresne).")
    doc.add_paragraph("- Règlement complet : https://www.mairie-latresne.fr/risques-inondations/")
