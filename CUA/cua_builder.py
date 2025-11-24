#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cua_builder_v6.py ‚Äî Builder principal pour g√©n√©ration du CUA DOCX
Nouveaut√©s v6 :
- Si 'reglementation' dans keep ‚Üí affichage uniquement de la r√©glementation
- Labels explicites : "Surface d'intersection" et "Pourcentage d'intersection"
"""

import argparse, os, logging, re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from CUA.cua_utils import (
    read_json, fmt_surface, fmt_pct, join_addr, parcels_label,
    build_footer_number, setup_doc, set_footer_num,
    add_first_article_title, add_article_title, add_paragraph, add_kv_table, add_objects_table,
    filter_intersections, filter_zonage_plu, equilibrer_pourcentages,
    add_annexes_section, ensure_page_space_for_article,
)

from CUA.cas_speciaux import appliquer_cas_speciaux
from CUA.cua_header import render_first_page_header, add_mayor_section_with_vu
from CUA.ppri_cua_module import analyser_ppri_corrige, generer_rapport_cua_avec_table


# ============================================================
# üÜï HELPER : D√©tection attribut "reglementation" dans keep
# ============================================================
def has_reglementation_in_keep(layer_key: str, catalogue: Dict[str, Any]) -> bool:
    """V√©rifie si 'reglementation' est dans les attributs keep du catalogue."""
    keep = catalogue.get(layer_key, {}).get("keep", [])
    return "reglementation" in keep


# ============================================================
# üÜï FONCTION D'AFFICHAGE PAR COUCHE (articles 3 √† 7)
# ============================================================
def render_layer_content(
    doc,
    layer: Dict[str, Any],
    layer_key: str,
    catalogue: Dict[str, Any],
    add_annexes_callback=None,
    force_table_mode=False  # ‚úÖ Nouveau param√®tre pour forcer le mode tableau
) -> None:
    """
    Affiche le contenu d'une couche selon la logique :
    - Si force_table_mode=True ‚Üí afficher tableau (zonage PLU)
    - Si 'reglementation' dans keep ‚Üí afficher uniquement la r√©glementation
    - Sinon ‚Üí afficher tableau des objets
    
    Args:
        doc: Document DOCX
        layer: Donn√©es de la couche (avec nom, surface_m2, pourcentage, objets)
        layer_key: Cl√© de la couche dans le catalogue
        catalogue: Catalogue complet des couches
        add_annexes_callback: Fonction pour ajouter des annexes (pour PLU)
        force_table_mode: Force le mode tableau m√™me si 'reglementation' dans keep
    """
    nom = layer.get("nom") or "Couche"
    surface_m2 = layer.get("surface_m2", 0)   # encore utile en interne si besoin
    pourcentage = layer.get("pourcentage", 0)
    objets = layer.get("objets") or []
    print(f"üîç DEBUG render_layer_content d√©but: layer_key={layer_key}, nb_objets={len(objets)}, force_table_mode={force_table_mode}")
    
    # Titre de la couche
    add_paragraph(doc, nom, bold=True)
    
    # ‚úÖ On n'affiche plus que le pourcentage, pas les m¬≤
    if pourcentage:
        add_paragraph(
            doc,
            f"Part de l'unit√© fonci√®re concern√©e : {fmt_pct(pourcentage)} de la surface cadastrale indicative."
        )
    else:
        # ‚úÖ V√©rifier le type de g√©om√©trie depuis le catalogue
        geom_type = catalogue.get(layer_key, {}).get("geom_type")
        
        if geom_type in ("lineaire", "ponctuelle"):
            add_paragraph(
                doc,
                f"Part de l'unit√© fonci√®re concern√©e : entit√©s {geom_type}s (pas de surface mesurable).",
                italic=True,
            )
        else:
            add_paragraph(
                doc,
                "Part de l'unit√© fonci√®re concern√©e : ‚Äî",
                italic=True,
            )
    
    # ============================================================
    # LOGIQUE CONDITIONNELLE
    # ============================================================
    # ‚úÖ Exception pour zonage PLU ou mode tableau standard
    if force_table_mode or not has_reglementation_in_keep(layer_key, catalogue):
        # ‚úÖ MODE TABLEAU
        reglements_annexes = []
        objets_pour_table = []
        
        # Fonction pour retirer les colonnes de surfaces (m¬≤) du tableau
        # ‚úÖ Ne supprimer que les colonnes de surface, PAS la r√©glementation
        def _strip_surface_keys(d):
            return {
                k: v
                for k, v in d.items()
                if not k.lower().startswith("surface")
                and not k.lower().endswith("_m2")
            }
        
        # ‚úÖ Pas de d√©doublonnage ici : d√©j√† fait dans filter_zonage_plu() pour l'article 3
        print(f"üîç DEBUG render_layer_content: {len(objets)} objet(s) √† traiter pour {layer_key}")
        for obj in objets:
            print(f"üîç Objet brut : {obj}")  # DEBUG
            # ‚úÖ Extraire la r√©glementation AVANT de traiter l'objet
            if "reglementation" in obj and obj["reglementation"]:
                reglements_annexes.append(obj["reglementation"])
            # ‚úÖ Cr√©er une copie sans r√©glementation pour le tableau
            obj_sans_regl = {k: v for k, v in obj.items() if k != "reglementation"}
            obj_sans_regl = _strip_surface_keys(obj_sans_regl)
            print(f"üîç Objet apr√®s _strip_surface_keys : {obj_sans_regl}")  # DEBUG
            objets_pour_table.append(obj_sans_regl)
        print(f"üîç DEBUG: {len(objets_pour_table)} objet(s) dans objets_pour_table apr√®s traitement")
        
        # Afficher le tableau si des objets existent
        if objets_pour_table:
            add_objects_table(doc, objets_pour_table)
        else:
            add_paragraph(doc, "Aucune information d√©taill√©e disponible.", italic=True)
        
        # Ajouter r√©glementations en annexe (sp√©cifique au zonage PLU)
        if reglements_annexes and add_annexes_callback:
            add_annexes_callback({
                "titre": f"R√®glement du PLU ‚Äì {nom}",
                "contenu": "\n\n".join(reglements_annexes)
            })
            add_paragraph(doc, "‚Üí Le texte complet du r√®glement du PLU est renvoy√© en annexe.", italic=True)
    
    else:
        # ‚úÖ MODE R√âGLEMENTATION UNIQUEMENT
        reglementations = []
        for obj in objets:
            if "reglementation" in obj and obj["reglementation"]:
                reglementations.append(str(obj["reglementation"]).strip())
        
        # ‚úÖ D√©dupliquer les r√©glementations avant affichage
        # Normaliser et garder uniquement les r√©glementations uniques
        reglementations_uniques = []
        reglementations_vues = set()
        
        for regl in reglementations:
            # Normaliser pour la comparaison (espaces multiples ‚Üí espace unique)
            normalized = re.sub(r'\s+', ' ', regl.strip()) if regl else ""
            if normalized and normalized not in reglementations_vues:
                reglementations_vues.add(normalized)
                reglementations_uniques.append(regl)  # Garder le texte original
        
        if reglementations_uniques:
            for reglement in reglementations_uniques:
                add_paragraph(doc, reglement)
        else:
            add_paragraph(doc, "Aucune r√©glementation sp√©cifique disponible.", italic=True)


# ====================== BUILD CUA DOC ======================

def build_cua_docx(
    cerfa_json: Dict[str, Any],
    intersections_json: Dict[str, Any],
    catalogue_json: Dict[str, Any],
    output_docx: str,
    *,
    wkt_path: Optional[str] = None,
    logo_first_page: Optional[str] = None,
    signature_logo: Optional[str] = None,
    qr_url: str = "https://www.kerelia.com/carte",
    plu_nom="PLU en vigueur",
    plu_date_appro="13/02/2017",
) -> None:

    meta = cerfa_json.get("data") or {}
    commune = (meta.get("commune_nom") or "‚Äî").upper()
    parcelles = parcels_label(meta.get("references_cadastrales") or [])
    terrain = join_addr(meta.get("adresse_terrain") or {})
    footer_num = build_footer_number(meta)
    ncu = meta.get("numero_cu") or "‚Äî"

    # --- Surface indicative : s√©curisation / cast en float ---
    raw_surface_total = meta.get("superficie_totale_m2")

    surface_total = None
    if raw_surface_total not in (None, "", " "):
        try:
            # Support format europ√©en (ex: "1 649,5")
            cleaned = str(raw_surface_total).replace(" ", "").replace("\u202f", "").replace(",", ".")
            surface_total = float(cleaned)
        except (TypeError, ValueError):
            raise ValueError(f"[CUA] superficie_totale_m2 invalide dans le CERFA: {raw_surface_total!r}")

    if surface_total is None:
        raise ValueError("[CUA] surface_indicative manquante ou invalide dans le CERFA.")

    inters = intersections_json or {}
    # ‚úÖ Surface indicative = surface du CERFA (pas la surface SIG)
    surface_indicative = surface_total
    
    # üîé Log : on plaque toutes les intersections sur la surface cadastrale CERFA
    logger = logging.getLogger("cua_builder")
    logger.info(
        f"üìè Surface indicative CERFA utilis√©e pour normaliser les intersections : {surface_indicative} m¬≤"
    )
    
    intersections_raw = inters.get("intersections") or {}
    
    # ‚úÖ Recalcul des surfaces finales indicatives √† partir des pourcentages SIG
    for key, layer in intersections_raw.items():
        pct_sig = layer.get("pct_sig", 0)
        # Surface indicative = pourcentage SIG * surface indicative CERFA
        layer["surface_m2"] = round((pct_sig / 100.0) * surface_indicative, 2)
        # Pourcentage final = pourcentage SIG (on le conserve tel quel)
        layer["pourcentage"] = pct_sig

    # Normalisation des surfaces et pourcentages (sans filtrage par seuil)
    # Utiliser surface_indicative pour les calculs
    intersections = filter_intersections(
        intersections_raw,
        catalogue_json,
        surface_indicative,
        min_pct=0.0  # ‚úÖ Garder toutes les couches, filtrer apr√®s sur objets
    )

    print(f"\nüîç DEBUG: Couches apr√®s filtrage:")
    for key in intersections.keys():
        print(f"   - {key}")

    # Initialisation du regroupement par article
    layers_by_article: Dict[str, List] = {}

    # Regroupement des couches selon leur article (AVANT cas sp√©ciaux)
    unknown_layers = []
    for key, layer in intersections.items():
        raw_art = catalogue_json.get(key, {}).get("article", None)
        
        # Normalisation de l'article en liste d'entiers
        if raw_art is None:
            article_list = None
        elif isinstance(raw_art, str) and "," in raw_art:
            # Cas "7, 5" ‚Üí [7, 5]
            article_list = [int(x.strip()) for x in raw_art.split(",") if x.strip().isdigit()]
        elif isinstance(raw_art, str):
            # Cas "3" ‚Üí [3]
            article_list = [int(raw_art.strip())] if raw_art.strip().isdigit() else None
        elif isinstance(raw_art, int):
            # Cas 3 ‚Üí [3]
            article_list = [raw_art]
        else:
            article_list = None
        
        # Tri dans les articles correspondants
        if article_list:
            for a in article_list:
                article_str = str(a)
                layers_by_article.setdefault(article_str, []).append((key, layer))
        else:
            unknown_layers.append(key)

    print(f"\nüîç DEBUG: layers_by_article['3'] = {layers_by_article.get('3')}")

    # Application des cas particuliers (apr√®s)
    # appliquer_cas_speciaux(intersections, layers_by_article)  # D√âSACTIV√â temporairement

    if unknown_layers:
        print("\n‚ö†Ô∏è  Les couches suivantes n'ont pas d'article d√©fini dans le catalogue :")
        for k in unknown_layers:
            print(f"   - {k}")

    # ‚úÖ Zonage PLU (Article 3) : garder les objets tels quels, pas de filtrage par seuil
    if layers_by_article.get("3"):
        for layer_key, layer_data in layers_by_article["3"]:
            # S√©curit√© : garder les objets tels quels, utiliser pct_sig comme pourcentage
            if layer_data.get("objets"):
                pct_sig = layer_data.get("pct_sig", 0)
                layer_data["pourcentage"] = pct_sig
        print(f"‚úÖ Zonage PLU : {len(layers_by_article['3'])} zone(s) conserv√©e(s)")

    # √âquilibrage des pourcentages dans chaque article
    for art, layer_tuples in layers_by_article.items():
        # Extraire les layers et leurs cl√©s pour √©quilibrage
        layers_data = [layer for _, layer in layer_tuples]
        layer_keys = [key for key, _ in layer_tuples]
        balanced_layers = equilibrer_pourcentages(layers_data, layer_keys=layer_keys, catalogue=catalogue_json)
        # Reconstituer les tuples
        layers_by_article[art] = [(layer_tuples[i][0], balanced_layers[i]) 
                                   for i in range(len(layer_tuples))]

    # Initialisation des annexes
    annexes = []

    # DOCX init
    doc = setup_doc()
    set_footer_num(doc, footer_num)

    # Header premi√®re page avec QR code
    render_first_page_header(
        doc,
        cerfa_json,
        logo_commune_path=logo_first_page,
        qr_url=qr_url,
        qr_logo_path=signature_logo
    )
    
    # Section "Le Maire" avec Vu et CERTIFIE
    add_mayor_section_with_vu(doc, cerfa_json, commune, plu_date_appro)

    # Article 1
    add_first_article_title(doc, "Article UN - Objet")
    add_paragraph(doc,
        "Les r√®gles d'urbanisme, la liste des taxes et participations d'urbanisme ainsi que "
        "les limitations administratives au droit de propri√©t√© applicables au terrain sont "
        "mentionn√©es aux articles 2 et suivants du pr√©sent certificat.\n\n"
        "Conform√©ment au quatri√®me alin√©a de l'article L. 410-1 du code de l'urbanisme, "
        "si une demande de permis de construire, d'am√©nager ou de d√©molir ou si une d√©claration "
        "pr√©alable est d√©pos√©e dans le d√©lai de dix-huit mois √† compter de la date du pr√©sent "
        "certificat d'urbanisme, les dispositions d'urbanisme, le r√©gime des taxes et participations "
        "d'urbanisme ainsi que les limitations administratives au droit de propri√©t√© tels qu'ils "
        "existaient √† cette date ne peuvent √™tre remis en cause √† l'exception des dispositions qui "
        "ont pour objet la pr√©servation de la s√©curit√© ou de la salubrit√© publique."
    )

    # Article 2
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article DEUX - Identification et localisation du terrain")
    add_kv_table(doc, [
        ("Commune", f"{meta.get('commune_nom') or '‚Äî'} ({meta.get('commune_insee') or '‚Äî'})"),
        ("Adresse / Localisation", terrain),
        ("R√©f√©rences cadastrales", parcelles),
        ("Surface indicative", (fmt_surface(surface_total) + " m¬≤") if surface_total else "‚Äî"),
        ("Document d'urbanisme opposable", f"{plu_nom} ‚Äî approuv√© le {plu_date_appro}")
    ])

    # Article 3 : Zonage
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article TROIS - Dispositions d'urbanisme (Zonage)")
    add_paragraph(doc,
        "Les occupations et utilisations du sol, ainsi que les r√®gles de constructibilit√©, "
        "sont d√©finies par le r√®glement du PLU. Ci-dessous, les th√©matiques majeures sont "
        "rappel√©es de mani√®re neutre avec renvoi aux articles sources (le texte du r√®glement fait foi)."
    )

    if layers_by_article.get("3"):
        for layer_key, layer_data in layers_by_article["3"]:
            if layer_data.get("objets"):  # ‚úÖ Ajout
                render_layer_content(
                    doc, 
                    layer_data, 
                    layer_key, 
                    catalogue_json,
                    add_annexes_callback=lambda annex: annexes.append(annex),
                    force_table_mode=True  # ‚úÖ Force mode tableau pour zonage PLU
                )
    else:
        add_paragraph(doc, "Aucune donn√©e de zonage disponible.", italic=True)

    # Article 4 : Servitudes d'utilit√© publique
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article QUATRE - Servitudes d'utilit√© publique (SUP)")
    
    if layers_by_article.get("4"):
        for layer_key, layer_data in layers_by_article["4"]:
            if layer_data.get("objets"):  # ‚úÖ Ajout
                render_layer_content(doc, layer_data, layer_key, catalogue_json)

    # Int√©gration automatique du PPRI PM1
    try:
        code_insee = meta.get("commune_insee") or "33234"
        print(f"üåä V√©rification du PPRI (PM1) pour l'unit√© fonci√®re (INSEE: {code_insee})‚Ä¶")

        if wkt_path and os.path.exists(wkt_path):
            with open(wkt_path, "r", encoding="utf-8") as f:
                geom_wkt = f.read().strip()
            resultats_ppri = analyser_ppri_corrige(geom_wkt=geom_wkt, code_insee=code_insee)
            print(f"‚úÖ WKT charg√© depuis : {wkt_path}")
        else:
            refs = meta.get("references_cadastrales", [])
            if refs:
                ref = refs[0]
                section = ref.get("section") or "AC"
                numero = ref.get("numero") or "0242"
            else:
                section, numero = "AC", "0242"
            print("‚ö†Ô∏è WKT non fourni, fallback cadastral.")
            resultats_ppri = analyser_ppri_corrige(section=section, numero=numero, code_insee=code_insee)

        if not resultats_ppri or not resultats_ppri.get("zones_avec_regles"):
            print("‚ÑπÔ∏è  Parcelle non concern√©e par le PPRI (aucune zone intersect√©e).")
        else:
            add_article_title(doc, "SUP PM1 ‚Äì Risques et Inondations (PPRI)")
            generer_rapport_cua_avec_table(doc, resultats_ppri)
            print("‚úÖ Rapport PPRI int√©gr√© dans la section PM1 avec tableau.")

    except Exception as e:
        print(f"‚ö†Ô∏è Erreur PPRI : {e}")

    add_paragraph(doc, "Avertissement : seuls les actes de servitudes publi√©s (et leurs annexes cartographiques) font foi.", italic=True)

    # Article 5 : Risques et protections environnementales
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article CINQ ‚Äì Risques et protections environnementales")
    
    if layers_by_article.get("5"):
        for layer_key, layer_data in layers_by_article["5"]:
            if layer_data.get("objets"):  # ‚úÖ Ajout
                render_layer_content(doc, layer_data, layer_key, catalogue_json)
    else:
        add_paragraph(doc, "Aucune donn√©e pertinente d√©tect√©e.", italic=True)

    # Article 6 : R√©seaux et √©quipements
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article SIX ‚Äì R√©seaux et √©quipements")
    
    if layers_by_article.get("6"):
        for layer_key, layer_data in layers_by_article["6"]:
            if layer_data.get("objets"):  # ‚úÖ Ajout
                render_layer_content(doc, layer_data, layer_key, catalogue_json)
    else:
        add_paragraph(doc, "Aucune donn√©e d'√©quipement disponible.", italic=True)

    # Article 7 : Informations utiles
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article SEPT ‚Äì Informations utiles")
    
    if layers_by_article.get("7"):
        for layer_key, layer_data in layers_by_article["7"]:
            if layer_data.get("objets"):  # ‚úÖ Ajout
                render_layer_content(doc, layer_data, layer_key, catalogue_json)
    else:
        add_paragraph(doc, "Aucune information compl√©mentaire d√©tect√©e.", italic=True)

    # Article 8 : Taxes et participations
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article HUIT ‚Äì Taxes et participations")
    add_paragraph(doc,
        "Les taxes suivantes pourront √™tre exig√©es √† compter de l'obtention d'un permis "
        "ou d'une d√©cision de non opposition √† une d√©claration pr√©alable."
    )
    add_kv_table(doc, [
        ("Taxe d'Am√©nagement", ""),
        ("Part communale - Taux en %", "5%"),
        ("Part d√©partementale - Taux en %", "2,5 %"),
        ("Redevance d'Arch√©ologie Pr√©ventive - Taux en %", "0,68 %")
    ])
    add_paragraph(doc, "Participations :", bold=True)
    add_paragraph(doc,
        "Les participations ci-dessous pourront √™tre exig√©es √† l'occasion d'un permis de construire "
        "ou d'une d√©cision de non opposition √† une d√©claration pr√©alable. Si tel est le cas elles "
        "seront mentionn√©es dans l'arr√™t√© de permis ou dans un arr√™t√© pris dans les deux mois "
        "suivant la date du permis tacite ou de la d√©cision de non opposition √† une d√©claration pr√©alable."
    )
    add_paragraph(doc, "Participations susceptibles d'√™tre exig√©es √† l'occasion de l'op√©ration :")
    add_paragraph(doc, "- contribution aux d√©penses de r√©alisation des √©quipements publics.")
    add_paragraph(doc, "- financement de branchements des √©quipements propres (article L332-15 du CU).")

    # Article 9 : Droit de pr√©emption
    ensure_page_space_for_article(doc)
    add_article_title(doc, "Article NEUF ‚Äì Droit de pr√©emption")
    
    has_dpu = bool(layers_by_article.get("9"))
    
    if has_dpu:
        add_paragraph(doc,
            "Le bien est situ√© dans un p√©rim√®tre de DPU (Droit de Pr√©emption Urbain) d√©limit√© "
            "au PLU ‚Äì annexe 6-1. Toute ali√©nation √† titre on√©reux est soumise √† DIA "
            "(D√©claration d'Intention d'Ali√©ner (C. urb. L211-1 s.). La commune dispose d'un "
            "d√©lai de 2 mois pour se prononcer, d√©lai suspendu en cas de demande unique de "
            "pi√®ces/visite (C. urb. L213-2). Silence = renonciation (C. urb. R213-4 s.). "
            "En cas de d√©saccord sur le prix, saisine du juge de l'expropriation (C. urb. L213-4)."
        )
        for layer_key, layer_data in layers_by_article["9"]:
            render_layer_content(doc, layer_data, layer_key, catalogue_json)
    else:
        add_paragraph(doc,
            "Le terrain n'est pas situ√© dans une zone de droit de pr√©emption. Aucune DIA "
            "(D√©claration d'Intention d'Ali√©ner) au titre du DPU (Droit de Pr√©emption Urbain) "
            "n'est requise."
        )

    # Signature
    doc.add_page_break()
    add_paragraph(doc, f"Fait √† {commune.title()}, le {datetime.now().strftime('%d/%m/%Y')}")
    add_paragraph(doc, "Le Maire,")
    if signature_logo and os.path.exists(signature_logo):
        try:
            doc.add_paragraph().add_run().add_picture(signature_logo, width=Cm(3))
        except Exception:
            pass
    
    # Informations finales
    doc.add_paragraph()
    p_info = doc.add_paragraph()
    r_info = p_info.add_run(
        "INFORMATIONS √Ä LIRE ATTENTIVEMENT\n\n"
        "Le (ou les) demandeur(s) peut contester la l√©galit√© de la d√©cision dans les deux mois qui suivent "
        "la date de sa notification. A cet effet il peut saisir le tribunal administratif territorialement "
        "comp√©tent d'un recours contentieux.\n\n"
        "Dur√©e de validit√© : Le certificat d'urbanisme a une dur√©e de validit√© de 18 mois. Il peut √™tre prorog√© "
        "par p√©riodes d'une ann√©e si les prescriptions d'urbanisme, les servitudes d'urbanisme de tous ordres "
        "et le r√©gime des taxes et participations n'ont pas √©volu√©. Vous pouvez pr√©senter une demande de prorogation "
        "en adressant une demande sur papier libre, accompagn√©e du certificat pour lequel vous demandez la prorogation "
        "au moins deux mois avant l'expiration du d√©lai de validit√©.\n\n"
        "A d√©faut de notification d'une d√©cision expresse portant prorogation du certificat d'urbanisme dans le d√©lai "
        "de deux mois suivant la r√©ception en mairie de la demande, le silence gard√© par l'autorit√© comp√©tente vaut "
        "prorogation du certificat d'urbanisme. La prorogation prend effet au terme de la validit√© de la d√©cision "
        "initiale (Art. R. 410-17-1)\n\n"
        "Effets du certificat d'urbanisme : le certificat d'urbanisme est un acte administratif d'information, qui "
        "constate le droit applicable en mentionnant les possibilit√©s d'utilisation de votre terrain et les diff√©rentes "
        "contraintes qui peuvent l'affecter. Il n'a pas valeur d'autorisation pour la r√©alisation des travaux ou d'une "
        "op√©ration projet√©e.\n\n"
        "Le certificat d'urbanisme cr√©e aussi des droits √† votre √©gard. Si vous d√©posez une demande d'autorisation "
        "(par exemple une demande de permis de construire) dans le d√©lai de validit√© du certificat, les nouvelles "
        "dispositions d'urbanisme ou un nouveau r√©gime de taxes ne pourront pas vous √™tre oppos√©es, sauf exceptions "
        "relatives √† la pr√©servation de la s√©curit√© ou de la salubrit√© publique.\n\n"
        "QR Code : Le QR code permet d'acc√©der √† une Carte interactive des r√®gles applicables (zonage, SUP, risques, "
        "prescriptions, obligations, informations). Affichage informatif ; en cas de divergence, les pi√®ces √©crites "
        "et le r√®glement en vigueur font foi. Cette solution vous est propos√©e par KERELIA, soci√©t√© immatricul√©e "
        "944 763 275 au R.C.S. de Bordeaux\n\n"
        "Modalit√©s de calcul des surfaces : La surface totale de l'unit√© fonci√®re mentionn√©e dans ce document est la contenance cadastrale, qui a une valeur purement indicative. De m√™me, les surfaces partielles correspondant √† l'intersection avec des zones r√©glementaires sont des estimations obtenues en appliquant un pourcentage de superposition cartographique (calcul√© par Syst√®me d'Information G√©ographique) √† ladite surface totale indicative. Par cons√©quent, toutes les surfaces mentionn√©es dans ce certificat sont communiqu√©es √† titre informatif et sont d√©pourvues de valeur juridique. Seul un arpentage ou un bornage r√©alis√© par un g√©om√®tre-expert peut garantir la surface r√©elle du terrain et de ses fractions."
    )
    r_info.font.size = Pt(8)

    # Ajout des annexes
    if annexes:
        add_annexes_section(doc, annexes)
        print(f"üìé {len(annexes)} annexes ajout√©es en fin de CUA (r√®glements PLU).")

    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    print(f"\n‚úÖ CUA DOCX g√©n√©r√© : {output_docx}")


# ====================== FONCTION IMPORTABLE ======================

def run_builder(
    cerfa_json,
    intersections_json,
    catalogue_json,
    output_path,
    wkt_path=None,
    logo_first_page=None,
    signature_logo=None,
    qr_url="https://www.kerelia.com/carte",
    plu_nom="PLU en vigueur",
    plu_date_appro="13/02/2017"
):
    """
    Fonction importable pour g√©n√©rer un CUA DOCX sans passer par subprocess.
    Reprend exactement la logique actuelle du CLI.
    
    Args:
        cerfa_json: Chemin vers le fichier JSON CERFA ou dict
        intersections_json: Chemin vers le fichier JSON intersections ou dict
        catalogue_json: Chemin vers le fichier JSON catalogue ou dict
        output_path: Chemin de sortie pour le fichier DOCX
        wkt_path: Chemin vers le fichier WKT (optionnel)
        logo_first_page: Chemin vers le logo premi√®re page (optionnel)
        signature_logo: Chemin vers le logo signature (optionnel)
        qr_url: URL du QR code
        plu_nom: Nom du PLU
        plu_date_appro: Date d'approbation du PLU
    """
    # Lecture des fichiers JSON si ce sont des chemins
    if isinstance(cerfa_json, str):
        cerfa = read_json(cerfa_json)
    else:
        cerfa = cerfa_json
    
    if isinstance(intersections_json, str):
        inters = read_json(intersections_json)
    else:
        inters = intersections_json
    
    if isinstance(catalogue_json, str):
        catalogue = read_json(catalogue_json)
    else:
        catalogue = catalogue_json

    build_cua_docx(
        cerfa, inters, catalogue, output_path,
        wkt_path=wkt_path,
        logo_first_page=logo_first_page,
        signature_logo=signature_logo,
        qr_url=qr_url,
        plu_nom=plu_nom,
        plu_date_appro=plu_date_appro,
    )


# ====================== CLI ======================

def main():
    ap = argparse.ArgumentParser(description="CUA Builder v6")
    ap.add_argument("--cerfa-json", required=True)
    ap.add_argument("--intersections-json", required=True)
    ap.add_argument("--wkt-path", help="Chemin WKT unit√© fonci√®re")
    ap.add_argument("--catalogue-json", required=True)
    ap.add_argument("--output", default="CUA_final.docx")
    ap.add_argument("--logo-first-page", default="")
    ap.add_argument("--signature-logo", default="")
    ap.add_argument("--qr-url", default="https://www.kerelia.com/carte")
    ap.add_argument("--plu-nom", default="PLU en vigueur")
    ap.add_argument("--plu-date-appro", default="13/02/2017")
    args = ap.parse_args()

    run_builder(
        cerfa_json=args.cerfa_json,
        intersections_json=args.intersections_json,
        catalogue_json=args.catalogue_json,
        output_path=args.output,
        wkt_path=args.wkt_path,
        logo_first_page=args.logo_first_page or None,
        signature_logo=args.signature_logo or None,
        qr_url=args.qr_url,
        plu_nom=args.plu_nom,
        plu_date_appro=args.plu_date_appro,
    )

if __name__ == "__main__":
    main()