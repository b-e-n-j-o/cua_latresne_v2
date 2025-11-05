# -*- coding: utf-8 -*-
"""
cua_header.py — En-tête CUA (1ʳᵉ page)
- Logo commune en haut
- Titres centrés
- Tableau récap à gauche + QR code avec logo Kerelia à droite
- Section "Le Maire" avec tous les "Vu" et "CERTIFIE"
"""

import os, io, datetime
from typing import Any, Tuple
from docx.document import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- Helpers données CERFA ----------
def _date_fr(iso: str | None) -> str:
    if not iso:
        return ""
    try:
        d = datetime.date.fromisoformat(iso[:10])
        return d.strftime("%d/%m/%Y")
    except Exception:
        return iso or ""

def _safe(x, default=""):
    return default if x in (None, "", []) else x

def _join_addr(ad) -> str:
    """Construit une adresse depuis dict ou string."""
    if not ad:
        return ""
    
    # Si string directe, retourner tel quel
    if isinstance(ad, str):
        return ad.strip()
    
    # Sinon traiter comme dict
    parts = []
    if ad.get("numero"): parts.append(str(ad["numero"]).strip())
    if ad.get("voie"): parts.append(str(ad["voie"]).strip())
    if ad.get("lieu_dit"): parts.append(str(ad["lieu_dit"]).strip())
    line1 = " ".join(parts).strip()
    line2 = " ".join([_safe(ad.get("code_postal")), _safe(ad.get("ville"))]).strip()
    return (line1 + (", " + line2 if line2 else "")).strip()

def _demandeur_block(cerfa: dict) -> Tuple[str, str]:
    d = (cerfa.get("data") or {}).get("demandeur") or {}
    who = (d.get("denomination") or " ".join([_safe(d.get("prenom")), _safe(d.get("nom"))]).strip()).strip()
    siret = _safe(d.get("siret"))
    who_fmt = (who.upper() + (f" (SIRET {siret})" if siret else ""))
    domicile = _join_addr(((cerfa.get("data") or {}).get("coord_demandeur") or {}).get("adresse") or {})
    return who_fmt, domicile

def _terrain_addr(cerfa: dict) -> str:
    return _join_addr(((cerfa.get("data") or {}).get("adresse_terrain") or {}))

def _parcelles_label(cerfa: dict) -> str:
    refs = ((cerfa.get("data") or {}).get("references_cadastrales") or [])
    if not refs:
        return "—"
    labels = []
    for r in refs:
        sec = (r.get("section") or "").strip().upper()
        num = (r.get("numero") or "").strip().zfill(4)
        if sec or num:
            labels.append(f"{sec} {num}".strip())
    return ", ".join(labels) or "—"

# ---------- Helpers mise en page ----------
def _emu_to_cm(v: int) -> float:
    return float(v) / 360000.0

def _content_width_cm(section) -> float:
    return _emu_to_cm(section.page_width - section.left_margin - section.right_margin)

# ---------- QR avec logo Kerelia au centre ----------
def _make_qr_png_bytes(text: str, logo_path: str | None = None, box_size: int = 10, border: int = 4) -> bytes:
    """
    Génère un QR code PNG avec logo Kerelia au centre sur fond blanc carré.
    Nécessite: qrcode, pillow
    """
    try:
        import qrcode
        from PIL import Image
        
        # QR avec correction d'erreur élevée (tolère logo au centre)
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=box_size,
            border=border
        )
        qr.add_data(text or "")
        qr.make(fit=True)
        
        qr_img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
        
        # Ajouter logo au centre avec fond blanc carré
        if logo_path and os.path.exists(logo_path):
            try:
                logo = Image.open(logo_path)
                if logo.mode != 'RGBA':
                    logo = logo.convert("RGBA")
                
                qr_width, qr_height = qr_img.size
                
                # Taille maximale du logo: 35% de la largeur du QR
                max_logo_dimension = int(qr_width * 0.35)
                logo_orig_width, logo_orig_height = logo.size
                
                # Redimensionner en conservant le ratio
                if logo_orig_width > logo_orig_height:
                    logo_resized_width = max_logo_dimension
                    logo_resized_height = int(max_logo_dimension * (logo_orig_height / logo_orig_width))
                else:
                    logo_resized_height = max_logo_dimension
                    logo_resized_width = int(max_logo_dimension * (logo_orig_width / logo_orig_height))
                
                logo_resized = logo.resize((logo_resized_width, logo_resized_height), Image.Resampling.LANCZOS)
                
                # Créer fond blanc carré
                margin = 1
                square_bg_side = max(logo_resized_width, logo_resized_height) + (margin * 1)
                background = Image.new('RGB', (square_bg_side, square_bg_side), 'white')
                
                # Centrer le logo sur le fond blanc
                logo_x_on_bg = (square_bg_side - logo_resized_width) // 2
                logo_y_on_bg = (square_bg_side - logo_resized_height) // 2
                background.paste(logo_resized, (logo_x_on_bg, logo_y_on_bg), logo_resized)
                
                # Centrer le fond blanc + logo sur le QR
                pos_x = (qr_width - square_bg_side) // 2
                pos_y = (qr_height - square_bg_side) // 2
                qr_img.paste(background, (pos_x, pos_y))
                
            except Exception:
                pass
        
        buf = io.BytesIO()
        qr_img.save(buf, format="PNG")
        return buf.getvalue()
        
    except Exception:
        # Fallback: carré blanc
        from PIL import Image
        img = Image.new("RGB", (400, 400), "white")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()

# ---------- Header première page ----------
def render_first_page_header(
    doc: Document,
    cerfa: dict,
    logo_commune_path: str | None,
    qr_url: str,
    qr_logo_path: str | None = None
) -> None:
    """
    Génère l'en-tête de première page:
    - Logo commune centré en haut
    - Titres centrés
    - Tableau infos à gauche (50%) + QR code à droite (50%)
    """
    data = cerfa.get("data") or {}
    commune = (data.get("commune_nom") or "").upper()
    
    # Logo commune centré
    if logo_commune_path and os.path.exists(logo_commune_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_logo.add_run().add_picture(logo_commune_path, width=Cm(4.0))
        except Exception:
            pass
    
    doc.add_paragraph()  # espace
    
    # Titres centrés
    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t1.add_run("CERTIFICAT D'URBANISME - Simple Information")
    r1.bold = True
    r1.font.size = Pt(16)
    
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("DÉLIVRÉ PAR LE MAIRE AU NOM DE LA COMMUNE")
    r2.bold = True
    r2.font.size = Pt(12)
    
    doc.add_paragraph()  # espace
    
    # Extraire données
    date_dep = _date_fr(data.get("date_depot"))
    who, domicile = _demandeur_block(cerfa)
    terrain = _terrain_addr(cerfa)
    num_cu = data.get("numero_cu") or ""
    
    # Table layout 1×2 (50/50)
    cw = _content_width_cm(doc.sections[0])
    left_w = (cw - 0.5) / 2.0
    right_w = (cw - 0.5) / 2.0
    
    layout = doc.add_table(rows=1, cols=2)
    layout.autofit = False
    layout.columns[0].width = Cm(left_w)
    layout.columns[1].width = Cm(right_w)
    
    # Gauche: tableau récap
    left_cell = layout.cell(0, 0)
    recap = left_cell.add_table(rows=5, cols=2)
    recap.style = "Table Grid"
    recap.autofit = False
    recap.columns[0].width = Cm(5.5)
    recap.columns[1].width = Cm(max(2.0, left_w - 5.7))
    
    rows_data = [
        ("Demande déposée le", date_dep),
        ("Par :", who),
        ("Demeurant à :", domicile),
        ("Sur un terrain sis à :", terrain),
        ("Lien vers la carte interactive :", "www.kerelia/labestcarte.com"),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        c0 = recap.cell(i, 0).paragraphs[0]
        c0.add_run(label).bold = False
        c1 = recap.cell(i, 1).paragraphs[0]
        r_val = c1.add_run(value or "—")
        if i < 4:  # Sauf URL
            r_val.bold = True
        # Pas de mise en forme spéciale pour l'URL
    
    # Droite: QR + numéro CU
    right_cell = layout.cell(0, 1)
    
    # Numéro CU en haut
    p_num = right_cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_num = p_num.add_run(f"N° CU {num_cu}")
    r_num.bold = True
    r_num.font.size = Pt(14)
    
    right_cell.add_paragraph()  # espace
    
    # QR code centré avec logo
    p_qr = right_cell.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_png = _make_qr_png_bytes(qr_url, logo_path=qr_logo_path)
    stream = io.BytesIO(qr_png)
    qr_size = min(right_w * 0.9, 6.0)
    p_qr.add_run().add_picture(stream, width=Cm(qr_size))

def add_mayor_section_with_vu(
    doc: Document,
    cerfa: dict,
    commune: str,
    plu_date_appro: str = "13/02/2017"
) -> None:
    """
    Section "Le Maire" avec tous les "Vu" et "CERTIFIE"
    """
    doc.add_page_break()
    
    data = cerfa.get("data") or {}
    date_dep = _date_fr(data.get("date_depot"))
    who, _ = _demandeur_block(cerfa)
    terrain = _terrain_addr(cerfa)
    parcelles = _parcelles_label(cerfa)
    num_cu = data.get("numero_cu") or ""
    
    # LE MAIRE (centré, gras, majuscules)
    p_maire = doc.add_paragraph()
    p_maire.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_maire = p_maire.add_run("LE MAIRE")
    r_maire.bold = True
    r_maire.font.size = Pt(14)
    
    doc.add_paragraph()  # espace
    
    # Vu la demande (dynamique avec données en gras)
    p_vu_demande = doc.add_paragraph()
    p_vu_demande.add_run(
        "Vu la demande d'un certificat d'urbanisme indiquant, en application de l'article L.410-1 a) "
        "du code de l'urbanisme, les dispositions d'urbanisme, les limitations administratives au droit "
        "de propriété et la liste des taxes et participations d'urbanisme applicables à un terrain situé à "
    )
    p_vu_demande.add_run(terrain or "—").bold = True
    p_vu_demande.add_run(" (cadastré ")
    p_vu_demande.add_run(parcelles).bold = True
    p_vu_demande.add_run("), présentée le ")
    p_vu_demande.add_run(date_dep or "—").bold = True
    p_vu_demande.add_run(" par ")
    p_vu_demande.add_run(who or "—").bold = True
    p_vu_demande.add_run(", et enregistrée par la mairie de ")
    p_vu_demande.add_run(commune.upper()).bold = True
    p_vu_demande.add_run(" sous le numéro ")
    p_vu_demande.add_run(num_cu.replace("-", "")).bold = True
    p_vu_demande.add_run(" ;")
    
    # Autres Vu (statiques)
    vu_texts = [
        "Vu le Code de l'Urbanisme et notamment ses articles L.410-1, R.410-1 et suivants ;",
        f"Vu le Plan Local d'urbanisme approuvé en date du {plu_date_appro} ;",
        "Vu le Plan de Prévention du risque naturel d'inondation (PPRNI) de l'Agglomération bordelaise - commune de Latresne, approuvé par arrêté préfectoral du 23 février 2022 ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'obligation de déclaration préalable lors de divisions foncières situées dans les zones naturelles et les zones agricoles sur le territoire de la commune ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'obligation de déclaration préalable lors de division du foncier bâti sur l'ensemble du territoire de la commune ;",
        "Vu la délibération du conseil municipal du 1er février 2024 instaurant l'autorisation préalable de travaux conduisant à la création de locaux à usage d'habitation dite « permis de diviser » sur l'ensemble du territoire de la commune ;"
    ]
    
    for vu_text in vu_texts:
        doc.add_paragraph(vu_text)
    
    doc.add_paragraph()  # espace
    
    # CERTIFIE (centré, gras, majuscules)
    p_certifie = doc.add_paragraph()
    p_certifie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_certifie = p_certifie.add_run("CERTIFIE :")
    r_certifie.bold = True
    r_certifie.font.size = Pt(14)
    
    doc.add_paragraph()  # espace

__all__ = [
    "render_first_page_header",
    "add_mayor_section_with_vu",
]