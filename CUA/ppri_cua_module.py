# -*- coding: utf-8 -*-
"""
ppri_cua_module.py
Module de g√©n√©ration de rapports PPRI pour documents CUA (Word, texte brut).
Fonctions d√©di√©es √† la production de contenu textuel et tableaux.
"""

import geopandas as gpd
from sqlalchemy import text
# --- Import du moteur SQL global d√©fini dans ppri_analyse_tolerance
try:
    from ppri_analyse_tolerance import engine as ENGINE_PPRI
    if ENGINE_PPRI is None:
        print("‚ö†Ô∏è ENGINE_PPRI est None ‚Äî tentative de reconstruction manuelle.")
        from dotenv import load_dotenv
        import os
        from sqlalchemy import create_engine
        load_dotenv()
        HOST = os.getenv("SUPABASE_HOST")
        DB = os.getenv("SUPABASE_DB")
        USER = os.getenv("SUPABASE_USER")
        PWD = os.getenv("SUPABASE_PASSWORD")
        PORT = os.getenv("SUPABASE_PORT", 5432)
        ENGINE_PPRI = create_engine(f"postgresql+psycopg2://{USER}:{PWD}@{HOST}:{PORT}/{DB}")
    else:
        print("‚úÖ ENGINE_PPRI import√© depuis ppri_analyse_tolerance.")
except Exception as e:
    print(f"‚ö†Ô∏è Impossible d'importer ENGINE_PPRI : {e}")
    ENGINE_PPRI = None


# --- Import du module altim√©trique IGN pour le paragraphe NGF g√©n√©ral
try:
    from cote_ngf import cote_ngf_parcelle
except ImportError:
    cote_ngf_parcelle = None
    print("‚ö†Ô∏è Module cote_ngf non disponible : paragraphe altim√©trique IGN d√©sactiv√©.")

# üîå Import de la connexion existante depuis ppri_analyse_tolerance
try:
    from ppri_analyse_tolerance import engine as ppri_engine
    engine = ppri_engine
    print("‚úÖ Engine PPRI import√© depuis ppri_analyse_tolerance.")
except Exception as e:
    engine = None
    print(f"‚ö†Ô∏è Impossible d'importer l'engine PPRI : {e}")


# ============================================================
# üìù G√âN√âRATION DE RAPPORTS POUR CUA
# ============================================================
def generer_rapport_cua(resultats):
    """
    G√©n√®re un rapport PPRI au format texte brut pr√™t √† ins√©rer dans le CUA (Article PM1).
    
    Args:
        resultats (dict): Dictionnaire retourn√© par analyser_ppri_corrige() ou analyser_ppri_tolerance()
                          Doit contenir:
                          - parcelle: {surface_m2, ...}
                          - zone_dominante: {nom, pourcentage, surface_m2, texte}
                          - zones_avec_regles: [{nom, pourcentage, surface_m2, texte}, ...]
                          - cas_multizone: bool
    
    Returns:
        str: Texte brut format√©, pr√™t √† ins√©rer dans le document CUA
    """
    
    dominante = resultats["zone_dominante"]
    zones = resultats["zones_avec_regles"]

    lignes = []
    lignes.append("Le PPRI combine zonage r√©glementaire et cotes de seuil (altitudes NGF de r√©f√©rence int√©grant le changement climatique √† horizon 2100).")
    lignes.append("")
    lignes.append("R√©partition surfacique :")
    lignes.append(f"La parcelle de {round(resultats['parcelle']['surface_m2'])} m¬≤ est r√©partie comme suit :")
    lignes.append("")

    for z in zones:
        lignes.append(f"- {z['nom']} : {z['pourcentage']} % ({z['surface_m2']} m¬≤)")

    if resultats["cas_multizone"]:
        lignes.append("")
        lignes.append("‚ö†Ô∏è Cas Multi-zones :")
        lignes.append("L'unit√© fonci√®re est concern√©e par plusieurs zones r√©glementaires. L'article A ‚Äì sec. IV ‚Äì ¬ß c du r√®glement du PPRI pr√©voit trois options d'application pour les projets :")
        lignes.append("1. Application stricte : adapter chaque partie du projet aux prescriptions de la zone qui la concerne.")
        lignes.append("2. R√®gle la plus restrictive : appliquer les prescriptions de la zone la plus contraignante √† l'ensemble du projet.")
        lignes.append("3. Implantation : implanter le projet enti√®rement dans la zone la moins expos√©e, si possible.")
        lignes.append("")
        lignes.append("Pour les b√¢timents existants √† cheval sur plusieurs zones (hors grenat et rouge non urbanis√©e), il est possible d'appliquer la r√©glementation de la zone la moins contraignante si quatre conditions sont remplies :")
        lignes.append("- la zone la moins contraignante est majoritaire en surface ;")
        lignes.append("- l'acc√®s aux parties en zone plus expos√©e se fait uniquement depuis la zone moins expos√©e ;")
        lignes.append("- pas d'augmentation de vuln√©rabilit√© (planchers au-dessus de la cote de seuil) ;")
        lignes.append("- la partie la plus expos√©e ne communique pas avec l'ext√©rieur sous la cote de seuil.")
        lignes.append("")

    # üîπ Cotes de seuil PPRI
    cotes = resultats.get("cotes_ngf", [])
    if cotes:
        min_cote = round(min(cotes), 2)
        max_cote = round(max(cotes), 2)
        if abs(max_cote - min_cote) < 0.05:
            lignes.append(f"Cote de seuil relev√©e : {min_cote} m NGF.")
        else:
            lignes.append(f"Cotes de seuil relev√©es : entre {min_cote} m et {max_cote} m NGF.")
        lignes.append("")

    # üîπ Paragraphe altim√©trique IGN
    if resultats.get("paragraphe_altitude"):
        lignes.append(resultats["paragraphe_altitude"])
        lignes.append("")

    lignes.append("Prescriptions d√©taill√©es :")
    lignes.append("Les textes suivants sont applicables aux diff√©rentes parties de la parcelle :")
    lignes.append("")

    for z in zones:
        lignes.append(f"{z['nom'].upper()} ({z['pourcentage']} %)")
        lignes.append("-" * 60)
        lignes.append(z["texte"].replace("\n", " ").strip())
        lignes.append("")

    lignes.append("Glossaire :")
    lignes.append("- Al√©a : probabilit√© qu'une inondation se produise avec une certaine intensit√©.")
    lignes.append("- Cote de seuil : hauteur minimale du premier plancher habitable pour √™tre hors d'eau.")
    lignes.append("- Changement de destination : modification de l'usage d'un b√¢timent (ex. commerce ‚Üí logement).")
    lignes.append("- Vuln√©rabilit√© : degr√© d'exposition des personnes et biens au risque.")
    lignes.append("")
    lignes.append("R√©f√©rences :")
    lignes.append("- Arr√™t√© pr√©fectoral du 23 f√©vrier 2022 (PPRI de Latresne).")
    lignes.append("- R√®glement complet : https://www.mairie-latresne.fr/risques-inondations/")
    lignes.append("")

    return "\n".join(lignes)


def generer_rapport_cua_structure(resultats):
    """
    Variante structur√©e retournant un dictionnaire pour plus de flexibilit√©.
    
    Args:
        resultats (dict): R√©sultats de l'analyse PPRI
    
    Returns:
        dict: {
            "introduction": str,
            "repartition": str,
            "multizone": str | None,
            "prescriptions": list[{zone: str, texte: str}],
            "glossaire": str,
            "references": str,
            "texte_complet": str
        }
    """
    
    zones = resultats["zones_avec_regles"]
    
    # Introduction
    introduction = "Le PPRI combine zonage r√©glementaire et cotes de seuil (altitudes NGF de r√©f√©rence int√©grant le changement climatique √† horizon 2100)."
    
    # R√©partition
    repartition_lignes = [
        "R√©partition surfacique :",
        f"La parcelle de {round(resultats['parcelle']['surface_m2'])} m¬≤ est r√©partie comme suit :",
        ""
    ]
    for z in zones:
        repartition_lignes.append(f"- {z['nom']} : {z['pourcentage']} % ({z['surface_m2']} m¬≤)")
    
    repartition = "\n".join(repartition_lignes)
    
    # Multi-zones
    multizone = None
    if resultats["cas_multizone"]:
        multizone_lignes = [
            "‚ö†Ô∏è Cas Multi-zones :",
            "L'unit√© fonci√®re est concern√©e par plusieurs zones r√©glementaires. L'article A ‚Äì sec. IV ‚Äì ¬ß c du r√®glement du PPRI pr√©voit trois options d'application pour les projets :",
            "1. Application stricte : adapter chaque partie du projet aux prescriptions de la zone qui la concerne.",
            "2. R√®gle la plus restrictive : appliquer les prescriptions de la zone la plus contraignante √† l'ensemble du projet.",
            "3. Implantation : implanter le projet enti√®rement dans la zone la moins expos√©e, si possible.",
            "",
            "Pour les b√¢timents existants √† cheval sur plusieurs zones (hors grenat et rouge non urbanis√©e), il est possible d'appliquer la r√©glementation de la zone la moins contraignante si quatre conditions sont remplies :",
            "- la zone la moins contraignante est majoritaire en surface ;",
            "- l'acc√®s aux parties en zone plus expos√©e se fait uniquement depuis la zone moins expos√©e ;",
            "- pas d'augmentation de vuln√©rabilit√© (planchers au-dessus de la cote de seuil) ;",
            "- la partie la plus expos√©e ne communique pas avec l'ext√©rieur sous la cote de seuil."
        ]
        multizone = "\n".join(multizone_lignes)
    
    # Prescriptions
    prescriptions = []
    for z in zones:
        prescriptions.append({
            "zone": f"{z['nom'].upper()} ({z['pourcentage']} %)",
            "texte": z["texte"].replace("\n", " ").strip()
        })
    
    # Glossaire
    glossaire = "\n".join([
        "Glossaire :",
        "- Al√©a : probabilit√© qu'une inondation se produise avec une certaine intensit√©.",
        "- Cote de seuil : hauteur minimale du premier plancher habitable pour √™tre hors d'eau.",
        "- Changement de destination : modification de l'usage d'un b√¢timent (ex. commerce ‚Üí logement).",
        "- Vuln√©rabilit√© : degr√© d'exposition des personnes et biens au risque."
    ])
    
    # R√©f√©rences
    references = "\n".join([
        "R√©f√©rences :",
        "- Arr√™t√© pr√©fectoral du 23 f√©vrier 2022 (PPRI de Latresne).",
        "- R√®glement complet : https://www.mairie-latresne.fr/risques-inondations/"
    ])
    
    # Texte complet
    texte_complet = generer_rapport_cua(resultats)
    
    return {
        "introduction": introduction,
        "repartition": repartition,
        "multizone": multizone,
        "prescriptions": prescriptions,
        "glossaire": glossaire,
        "references": references,
        "texte_complet": texte_complet
    }


# ============================================================
# üß© WRAPPER COMPATIBILIT√â (ancien nom de fonction)
# ============================================================
def analyser_ppri_corrige(geom_wkt=None, section=None, numero=None, code_insee=None, engine=None, **kwargs):
    """
    Wrapper r√©trocompatible pour cua_builder.py
    ‚Üí ex√©cute analyser_ppri_tolerance() et ajoute :
       - cotes de seuil PPRI (PostGIS)
       - paragraphe altim√©trique IGN (API Altim√©trie)
    
    Accepte d√©sormais :
    - geom_wkt : g√©om√©trie WKT directe (prioritaire)
    - section + numero : identifiants parcellaires (fallback)
    """
    print("‚öôÔ∏è [Compatibilit√©] Appel de analyser_ppri_corrige() redirig√© vers analyser_ppri_tolerance()")
    from ppri_analyse_tolerance import analyser_ppri_tolerance, engine as ppri_engine
    
    # ‚úÖ Utiliser ENGINE_PPRI si engine n'est pas fourni
    if engine is None:
        engine = ENGINE_PPRI or ppri_engine
    
    if engine is None:
        raise RuntimeError("‚ùå Moteur SQL introuvable pour l'analyse PPRI")

    # --- Ex√©cution de l'analyse de base
    resultats = analyser_ppri_tolerance(
        geom_wkt=geom_wkt,
        section=section,
        numero=numero,
        code_insee=code_insee,
        ppri_table="latresne.pm1_detaillee_gironde",
        engine=engine  # ‚úÖ passage explicite
    )

    if not resultats or "zones" not in resultats:
        return {}

    # Construction simplifi√©e pour compatibilit√© avec cua_builder
    cons = resultats["zones"]["ppri_conservees"]
    surface_parcelle = resultats["parcelle"]["surface"]
    
    zones_liste = []
    for _, row in cons.iterrows():
        surface_zone = row.geometry.area
        zones_liste.append({
            "nom": str(row.get("codezone", "Zone")),
            "pourcentage": round((surface_zone / surface_parcelle) * 100, 2),
            "surface_m2": round(surface_zone, 2),
            "texte": str(row.get("reglementation", "")).strip()
        })

    surface_totale = sum(z["surface_m2"] for z in zones_liste)
    zone_dominante = max(zones_liste, key=lambda z: z["surface_m2"], default={})

    # ============================================================
    # üßÆ √âTAPE 1 ‚Äî COTES DE SEUIL PPRI (colonne codezone = valeur NGF)
    # ============================================================
    cotes_ngf = []
    try:
        print("üåä Recherche des cotes de seuil PPRI (colonne codezone)...")
        sql_cote = """
            WITH parcelle AS (
                SELECT ST_GeomFromText(:wkt, 2154) AS geom
            )
            SELECT c.codezone AS cote_ngf
            FROM latresne.cote_de_seuil_ppri c, parcelle p
            WHERE ST_Intersects(c.geom_2154, p.geom);
        """
        with engine.connect() as conn:
            res = conn.execute(text(sql_cote), {"wkt": geom_wkt})
            cotes_ngf = [float(r[0]) for r in res.fetchall() if r[0] is not None]

        if cotes_ngf:
            print(f"‚úÖ {len(cotes_ngf)} cotes de seuil trouv√©es : {cotes_ngf}")
        else:
            print("‚ö†Ô∏è Aucune cote de seuil trouv√©e pour cette unit√© fonci√®re.")
    except Exception as e:
        print(f"‚ö†Ô∏è Erreur chargement cotes de seuil PPRI : {e}")

    # ============================================================
    # üß≠ √âTAPE 2 ‚Äî COTE ALTIM√âTRIQUE IGN (API Altim√©trie)
    # ============================================================
    paragraphe_altitude = ""
    if geom_wkt and cote_ngf_parcelle:
        try:
            print("üì° Calcul altim√©trique global de l'unit√© fonci√®re (API IGN)...")
            paragraphe_altitude = cote_ngf_parcelle(geom_wkt)
            print("‚úÖ Paragraphe altim√©trique IGN int√©gr√©.")
        except Exception as e:
            print(f"‚ö†Ô∏è Impossible de g√©n√©rer le paragraphe altim√©trique IGN : {e}")
    else:
        print("‚ÑπÔ∏è Calcul altim√©trique IGN d√©sactiv√© (g√©om√©trie WKT manquante).")

    # ============================================================
    # üîß STRUCTURE FINALIS√âE POUR LE BUILDER
    # ============================================================
    resultats_cua = {
        "parcelle": {"surface_m2": int(round(surface_totale))},
        "zone_dominante": zone_dominante,
        "zones_avec_regles": zones_liste,
        "cas_multizone": len(zones_liste) > 1,
        "cotes_ngf": cotes_ngf,
        "paragraphe_altitude": paragraphe_altitude
    }

    print(f"‚úÖ Zones distinctes : {len(zones_liste)}")
    print(f"‚úÖ Zone dominante : {zone_dominante.get('nom')} ({zone_dominante.get('pourcentage', 0)}%)")
    return resultats_cua


# ============================================================
# üß© WRAPPER COMPATIBILIT√â POUR generer_rapport_cua_avec_table
# ============================================================
def generer_rapport_cua_avec_table(doc, resultats):
    """
    Wrapper de compatibilit√© pour cua_builder.py.
    G√©n√®re un tableau PPRI synth√©tique dans le document DOCX.
    Utilise la logique de generer_rapport_cua() pour le texte principal.
    """
    from docx.shared import Pt
    from docx import Document

    # ‚úÖ Structure simplifi√©e compatible avec le CUA builder
    zones = resultats.get("zones_avec_regles", [])
    dominante = resultats.get("zone_dominante", {})
    surface_totale = int(round(resultats.get("parcelle", {}).get("surface_m2", 0)))

    # Titre et introduction
    doc.add_paragraph("Le PPRI combine zonage r√©glementaire et cotes de seuil (altitudes NGF de r√©f√©rence int√©grant le changement climatique √† horizon 2100).")
    doc.add_paragraph(f"La zone dominante est {dominante.get('nom', 'inconnue')} ({dominante.get('pourcentage', 0)} % de la parcelle).")
    doc.add_paragraph(f"La parcelle de {surface_totale} m¬≤ est r√©partie comme suit :")

    # Cr√©ation du tableau
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Zone").bold = True
    hdr_cells[1].paragraphs[0].add_run("Pourcentage").bold = True
    hdr_cells[2].paragraphs[0].add_run("Surface (m¬≤)").bold = True

    for z in zones:
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(z["nom"])
        row_cells[1].paragraphs[0].add_run(f"{z['pourcentage']} %")
        row_cells[2].paragraphs[0].add_run(str(int(round(z["surface_m2"], 0))))

    doc.add_paragraph("")  # espacement

    # üîπ Cotes de seuil PPRI (si dispo)
    cotes = resultats.get("cotes_ngf", [])
    if cotes:
        min_cote = round(min(cotes), 2)
        max_cote = round(max(cotes), 2)
        if abs(max_cote - min_cote) < 0.05:
            doc.add_paragraph(f"Cote de seuil relev√©e : {min_cote} m NGF.")
        else:
            doc.add_paragraph(f"Cotes de seuil relev√©es : entre {min_cote} m et {max_cote} m NGF.")
        doc.add_paragraph("")

    # üîπ Paragraphe altim√©trique IGN (si dispo)
    if resultats.get("paragraphe_altitude"):
        doc.add_paragraph(resultats["paragraphe_altitude"])
        doc.add_paragraph("")

    # üîπ Prescriptions r√©glementaires
    doc.add_paragraph("Prescriptions d√©taill√©es :", style="List Bullet")
    for z in zones:
        doc.add_paragraph(f"{z['nom'].upper()} ({z['pourcentage']} %) :", style="List Bullet")
        doc.add_paragraph(z["texte"].replace("\n", " ").strip())

    doc.add_paragraph("")
    doc.add_paragraph("R√©f√©rences :")
    doc.add_paragraph("- Arr√™t√© pr√©fectoral du 23 f√©vrier 2022 (PPRI de Latresne).")
    doc.add_paragraph("- R√®glement complet : https://www.mairie-latresne.fr/risques-inondations/")
