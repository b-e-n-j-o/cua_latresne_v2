# cua_utils.py
# -*- coding: utf-8 -*-
"""
Utilitaires communs pour le builder CUA (v4)
--------------------------------------------
- Formatage (dates, surfaces, %), lecture JSON
- Helpers DOCX (styles, tableaux, encadrés)
- Filtrage des intersections : supprime UNIQUEMENT les entités < min_pct
"""

import json, os
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ========================== FORMATAGE ==========================

def read_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def date_fr(d: Optional[str]) -> str:
    if not d:
        return "—"
    try:
        dt = datetime.strptime(d[:10], "%Y-%m-%d")
    except Exception:
        return d
    mois = [
        "janvier","février","mars","avril","mai","juin",
        "juillet","août","septembre","octobre","novembre","décembre"
    ]
    return f"{dt.day} {mois[dt.month-1]} {dt.year}"

def fmt_surface(val: Optional[float]) -> str:
    """Arrondit les surfaces au m² entier (sans décimales)."""
    if val is None:
        return "—"
    try:
        return str(int(round(float(val), 0)))
    except Exception:
        return str(val)

def fmt_pct(val: Optional[float]) -> str:
    """
    Formate les pourcentages avec deux décimales max (0.01 %).
    Corrige les dépassements de 100 % dus aux arrondis.
    Si >= 99%, affiche 100% pour une meilleure lisibilité.
    """
    if val is None:
        return "—"
    try:
        p = round(float(val), 2)
        # Si >= 99%, on affiche 100%
        if p >= 99.0:
            p = 100.00
        # Borne supérieure à 100% (par sécurité)
        if p > 100:
            p = 100.00
        return f"{p:.2f} %"
    except Exception:
        return str(val)

def join_addr(addr: Dict[str, Any]) -> str:
    if not addr:
        return "—"
    parts = []
    if addr.get("numero"): parts.append(str(addr["numero"]).strip())
    if addr.get("voie"): parts.append(str(addr["voie"]).strip())
    if addr.get("lieu_dit"): parts.append(str(addr["lieu_dit"]).strip())
    tail = []
    if addr.get("code_postal"): tail.append(str(addr["code_postal"]).strip())
    if addr.get("ville"): tail.append(str(addr["ville"]).strip())
    s1 = " ".join([p for p in parts if p])
    s2 = " ".join([t for t in tail if t])
    return ", ".join([p for p in [s1, s2] if p]) or "—"

def parcels_label(refs: List[Dict[str, Any]]) -> str:
    if not refs:
        return "—"
    lab = []
    for r in refs:
        sec = (r.get("section") or "").strip().upper()
        num = (r.get("numero") or "").strip().zfill(4)
        if sec or num:
            lab.append(f"{sec} {num}".strip())
    return ", ".join(lab) or "—"

def build_footer_number(meta: Dict[str, Any]) -> str:
    ncu = (meta.get("numero_cu") or "").strip()
    if ncu:
        return f"CU — {ncu}"
    hdr = meta.get("header_cu") or {}
    dep = (hdr.get("dept") or "").strip()
    com = (hdr.get("commune_code") or "").strip()
    an = (hdr.get("annee") or "").strip()
    nd = (hdr.get("numero_dossier") or "").strip()
    if dep and com and an and nd:
        return f"CU — {dep}-{com}-{an}-{nd}"
    return "Certificat d'urbanisme"


# ========================== DOCX HELPERS ==========================

ARTICLE_SPACE_AFTER_PT = 14

def setup_doc() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2); s.right_margin = Cm(2)
        s.different_first_page_header_footer = True
    return doc

def set_footer_num(doc: Document, text: str):
    for s in doc.sections:
        p = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.clear()
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9)

def add_title(doc: Document, title: str):
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def maybe_add_page_break(doc: Document, threshold_ratio: float = 0.5):
    """
    Force un saut de page si l'espace restant sur la page actuelle
    est inférieur à threshold_ratio (par défaut 0.5 = moitié de page).
    """
    section = doc.sections[-1]
    page_height = section.page_height.cm - (section.top_margin.cm + section.bottom_margin.cm)
    # Approximation de la hauteur de contenu : nombre de paragraphes * 0.6 cm
    content_height = sum(0.6 for p in doc.paragraphs if p.text.strip())
    used_ratio = min(content_height / page_height, 1.0)
    remaining_ratio = 1 - used_ratio

    if remaining_ratio < threshold_ratio:
        doc.add_page_break()

def _set_widow_orphan_control(paragraph):
    """Active le contrôle des veuves/orphelins et empêche les coupures de page."""
    pPr = paragraph._element.get_or_add_pPr()
    # widowControl : empêche les lignes orphelines
    widowControl = OxmlElement('w:widowControl')
    widowControl.set(qn('w:val'), '1')
    pPr.append(widowControl)
    # keepLines : garde toutes les lignes du paragraphe ensemble
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)

def add_first_article_title(doc: Document, title_text: str):
    """Ajoute le titre du premier article (sans espacement avant)."""
    p = doc.add_paragraph()
    r = p.add_run(title_text.upper()); r.bold = True; r.font.size = Pt(12)
    # 2 sauts de ligne après le titre (avant le contenu)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.keep_with_next = True
    # Empêche le titre d'être séparé du contenu
    _set_widow_orphan_control(p)
    return p

def add_article_title(doc: Document, title_text: str):
    """
    Ajoute le titre d'un article avec espacements (3 lignes avant, 2 après).
    Les propriétés widowControl et keep_with_next empêchent automatiquement
    qu'un titre soit seul en bas de page (Word gère cela intelligemment).
    """
    p = doc.add_paragraph()
    r = p.add_run(title_text.upper()); r.bold = True; r.font.size = Pt(12)
    # 3 sauts de ligne avant l'article (entre 2 articles)
    p.paragraph_format.space_before = Pt(36)
    # 2 sauts de ligne après le titre (avant le contenu)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.keep_with_next = True
    # Empêche le titre d'être séparé du contenu
    _set_widow_orphan_control(p)
    return p

def add_paragraph(doc: Document, text: str, *, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(ARTICLE_SPACE_AFTER_PT//2)
    return p

def add_kv_table(doc: Document, rows: List[Tuple[str,str]]):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid"
    for i, (k,v) in enumerate(rows):
        p0 = t.cell(i,0).paragraphs[0]; p0.add_run(k)
        p1 = t.cell(i,1).paragraphs[0]; r1 = p1.add_run(v or "—"); r1.bold = True
        for par in (p0, p1):
            par.space_after = Pt(2); par.space_before = Pt(2)
    doc.add_paragraph("")

def add_reglementation_block(doc: Document, texte: str):
    if not (texte and str(texte).strip()): 
        return
    p = doc.add_paragraph()
    shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
    p._element.get_or_add_pPr().append(shading_elm)
    p.add_run("Réglementation :").bold = True
    p.add_run("\n" + str(texte).strip())
    p.paragraph_format.space_after = Pt(6)

def add_objects_table(doc: Document, objets: List[Dict[str, Any]]):
    """
    Ajoute un tableau d'objets au document avec labels améliorés.
    Si tous les objets ont une surface nulle, ajoute une note explicative.
    """
    if not objets:
        return
    
    # Vérifie si tous les objets sont à surface nulle
    all_zero_surface = all(
        float(o.get("surface_inter_m2") or 0) == 0 for o in objets
    )

    # ✅ Cas 1 : uniquement des entités linéaires / ponctuelles (surface = 0)
    if all_zero_surface:
        add_paragraph(
            doc,
            "Étant donné qu'il s'agit d'entités linéaires ou ponctuelles, "
            "cet élément n'a pas de surface d'intersection mesurable, "
            "mais il est présent sur l'unité foncière.",
            italic=True,
        )
        # S'il existe des réglementations associées, on les affiche ensuite
        for o in objets:
            if o.get("reglementation"):
                add_reglementation_block(doc, o["reglementation"])
        doc.add_paragraph("")
        return

    # ✅ Cas 2 : entités surfaciques classiques (au moins une surface > 0)
    COLUMN_LABELS = {
        "surface_inter_m2": "Surface (m²)",
        "pourcentage_inter": "Pourcentage (%)",
        "zonage_reglement": "Zonage",
        "libelle": "Libellé",
        "categorie": "Catégorie",
        "type": "Type",
        "nom": "Nom"
    }
    
    all_keys = set()
    for obj in objets:
        all_keys.update(obj.keys())
    
    ignore_patterns = ["id", "uuid", "gid", "fid", "globalid", "geom", "reglementation"]
    keys = [k for k in sorted(all_keys) 
            if not any(pat in k.lower() for pat in ignore_patterns)]
    
    if not keys:
        return
    
    table = doc.add_table(rows=1 + len(objets), cols=len(keys))
    table.style = "Table Grid"
    
    for i, key in enumerate(keys):
        label = COLUMN_LABELS.get(key, key.replace("_", " ").title())
        cell = table.cell(0, i)
        cell.paragraphs[0].add_run(label).bold = True
    
    for row_idx, obj in enumerate(objets, start=1):
        for col_idx, key in enumerate(keys):
            val = obj.get(key, "")
            if val is None or val == "":
                display_val = "—"
            elif "surface" in key.lower() and "m2" in key.lower():
                display_val = str(int(round(float(val), 0)))
            elif "pourcentage" in key.lower():
                pct_val = round(float(val), 2)
                display_val = "100.00 %" if pct_val >= 99.0 else f"{pct_val:.2f} %"
            else:
                display_val = str(val)
            table.cell(row_idx, col_idx).text = display_val
    
    # Bloc(s) de réglementation en dessous du tableau
    for o in objets:
        if o.get("reglementation"):
            add_reglementation_block(doc, o["reglementation"])
    
    doc.add_paragraph("")


# ========================== FILTRAGE (ENTITÉS SEULEMENT) ==========================

def filter_intersections(
    intersections: Dict[str, Any],
    parcelle_surface: Optional[float],
    min_pct: float = 0.5
) -> Dict[str, Any]:
    """
    Filtrage et normalisation des intersections.

    ⚙️ Règles :
      - Garde TOUTES les couches ayant au moins un objet, même si surface_m2 = 0.
      - Garde aussi les couches avec surface > 0 (classique).
      - Supprime seulement les couches vides (aucun objet et aucune surface).
      - Arrondit surfaces et pourcentages.
    """
    if not intersections:
        return {}

    total = float(parcelle_surface or 0.0)
    normalized: Dict[str, Any] = {}

    for key, layer in intersections.items():
        objets = list(layer.get("objets") or [])
        normalized_objs: List[Dict[str, Any]] = []
        surf_cumul = 0.0

        for o in objets:
            s = float(o.get("surface_inter_m2") or 0.0)
            pct_obj = (s / total * 100.0) if total > 0 else 0.0
            o["surface_inter_m2"] = int(round(s))
            o["pourcentage_inter"] = round(pct_obj, 2)
            normalized_objs.append(o)
            surf_cumul += s

        # Mise à jour des valeurs normalisées
        new_layer = dict(layer)
        new_layer["objets"] = normalized_objs
        new_layer["surface_m2"] = int(round(surf_cumul))
        pct_layer = (surf_cumul / total * 100.0) if total > 0 else 0.0
        new_layer["pourcentage"] = round(pct_layer, 2)

        # ✅ Nouvelle logique de conservation :
        # Garder la couche si :
        # - elle contient au moins 1 objet (même surface = 0)
        # - OU elle a une surface cumulée > 0
        if normalized_objs or surf_cumul > 0:
            normalized[key] = new_layer

    return normalized

# ====================== FILTRAGE SPÉCIFIQUE ZONAGE PLU ======================

def filter_zonage_plu(layer: Dict[str, Any], parcelle_surface: float, min_pct: float = 1.0) -> Dict[str, Any]:
    """
    Filtre spécifique pour le zonage PLU :
      - Supprime les objets (zones) < min_pct (par défaut 1%)
      - Rééquilibre les pourcentages des zones conservées pour totaliser 100%
      - Recalcule la surface totale de la couche
    
    Args:
        layer: Couche de zonage PLU avec objets
        parcelle_surface: Surface totale de la parcelle
        min_pct: Seuil minimum en % (défaut 1.0%)
    
    Returns:
        Couche filtrée et rééquilibrée
    """
    objets = list(layer.get("objets") or [])
    if not objets or parcelle_surface <= 0:
        return layer
    
    kept_objs = []
    surf_cumul = 0.0
    
    # Filtrer les zones < 1%
    for obj in objets:
        s = float(obj.get("surface_inter_m2") or 0.0)
        pct_obj = (s / parcelle_surface) * 100.0
        
        if pct_obj >= min_pct:
            kept_objs.append(obj)
            surf_cumul += s
    
    # Si aucune zone conservée, retourner la couche vide
    if not kept_objs:
        return None
    
    # Rééquilibrer les pourcentages pour totaliser 100%
    # On considère que les zones conservées représentent 100% du zonage
    for obj in kept_objs:
        s = float(obj.get("surface_inter_m2") or 0.0)
        # Nouveau pourcentage = (surface zone / surface totale des zones conservées) * 100
        obj["pourcentage_inter"] = round((s / surf_cumul) * 100.0, 2) if surf_cumul > 0 else 0.0
        obj["surface_inter_m2"] = int(round(s))
    
    # Mettre à jour la couche
    new_layer = dict(layer)
    new_layer["objets"] = kept_objs
    new_layer["surface_m2"] = int(round(surf_cumul))
    new_layer["pourcentage"] = 100.0  # Les zones conservées = 100% du zonage
    
    return new_layer


# ====================== ÉQUILIBRAGE DES POURCENTAGES ======================

def equilibrer_pourcentages(zones, key="pourcentage", precision=2):
    """
    Ajuste les pourcentages d'une liste de zones pour que la somme fasse exactement 100.00 %.
    - zones : liste de dicts contenant un champ 'pourcentage'
    - key : nom de la clé contenant la valeur à équilibrer
    - precision : nombre de décimales à conserver (par défaut 2)
    Retourne la liste corrigée.
    """
    if not zones or not isinstance(zones, list):
        return zones

    total = round(sum(float(z.get(key, 0) or 0) for z in zones), precision)
    diff = round(100.0 - total, precision)

    # Si la somme est déjà correcte, rien à faire
    if abs(diff) < (10 ** -precision):
        return zones

    # Trouver la zone dominante (plus grande surface)
    idx_max = max(range(len(zones)), key=lambda i: zones[i].get(key, 0) or 0)

    # Appliquer la correction sur la zone dominante
    zones[idx_max][key] = round((zones[idx_max].get(key, 0) or 0) + diff, precision)

    # Garantir bornes 0–100 après correction
    for z in zones:
        z[key] = min(max(round(float(z.get(key, 0)), precision), 0.0), 100.0)

    return zones


# ====================== ANNEXES ======================

def add_annexes_section(doc: Document, annexes: List[Dict[str, str]]):
    """
    Ajoute une section 'Annexes' à la fin du document DOCX.
    Chaque annexe est un dict avec : {'titre': str, 'contenu': str}
    """
    if not annexes:
        return

    doc.add_page_break()
    add_article_title(doc, "ANNEXES")

    for ann in annexes:
        titre = ann.get("titre", "Annexe")
        contenu = ann.get("contenu", "").strip()
        if not contenu:
            continue
        add_paragraph(doc, titre, bold=True)
        add_paragraph(doc, contenu)

def ensure_page_space_for_article(doc, threshold_ratio: float = 0.5):
    """
    Vérifie l'espace restant sur la page actuelle et force un saut de page
    si le titre du prochain article risquerait d'être trop bas (en dessous du seuil).
    threshold_ratio = 0.5 (par défaut = moitié de page)
    """
    try:
        section = doc.sections[-1]
        page_height = section.page_height.cm - (section.top_margin.cm + section.bottom_margin.cm)
        # Estimation de la hauteur déjà utilisée : 0.6 cm par paragraphe non vide
        used_height = sum(0.6 for p in doc.paragraphs if p.text.strip())
        used_ratio = min(used_height / page_height, 1.0)
        remaining_ratio = 1 - used_ratio

        if remaining_ratio < threshold_ratio:
            doc.add_page_break()
    except Exception as e:
        print(f"⚠️ Erreur lors de la vérification d'espace page : {e}")



