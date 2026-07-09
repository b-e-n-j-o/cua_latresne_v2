# -*- coding: utf-8 -*-
"""
builder.py — Génération du Certificat d'Urbanisme (DOCX) pour Argelès-sur-Mer.

Consomme le rapport d'intersections (sortie de intersections.py) : pour chaque
couche concernée, lit la réglementation déjà taguée en base (obj["reglementation"])
et l'écrit dans la bonne section du document. Aucune logique réglementaire ici —
le texte vit en base, le builder ne fait que router et mettre en forme.

Architecture : une fonction par section, le builder itère sur SECTIONS.
- Changer de commune  → nouveau CommuneConfig.
- Logo commune en haut à droite de la 1ʳᵉ page (n° dossier / pagination : à brancher).
"""

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from api.cuas.argeles.docx_utils import (
        add_markdown_block,
        add_para_with_link,
        parse_markdown_inline,
    )
except ImportError:
    from docx_utils import add_markdown_block, add_para_with_link, parse_markdown_inline

try:
    from api.modules_communs.intersection_partielle import (
        est_multi_entites,
        note_pour_objet,
        note_pour_servitude,
        texte_note_intersection_partielle,
    )
except ImportError:
    def est_multi_entites(objets):
        return len(objets or []) > 1

    def texte_note_intersection_partielle(pct, *, multi_entites=False, enabled=True):
        return None

    def note_pour_objet(obj, *, multi_entites=False, enabled=True):
        return None

    def note_pour_servitude(servitude, *, multi_entites_couche=False, enabled=True, surface_sig=0.0):
        return None

_ARGELES_DIR = Path(__file__).resolve().parent
LOGO_COMMUNE_PATH = _ARGELES_DIR / "logos" / "argeles.png"


# ============================================================
# CONFIG COMMUNE
# ============================================================
@dataclass
class CommuneConfig:
    nom: str = "ARGELÈS-SUR-MER"
    code_insee: str = "66008"
    departement: str = "Pyrénées-Orientales"

    plu_mention: str = ("approuvé le 20/04/2017, révisé le 10/03/2022, "
                        "modifié le 14/12/2023 et le 30/10/2025")
    pprn_mention: str = ("approuvé par arrêté préfectoral du 25/11/2008, "
                         "modifié par arrêté préfectoral du 29/05/2017")
    pgri_mention: str = ("porter à connaissance relatif aux règles de gestion du risque "
                         "d'inondation (PGRI) en date du 11/07/2019")
    geoportail_url: str = "https://data.geopf.fr/annexes/gpu/documents/DU_66008/dd4c8deab39aa04c8938e88dd2337dba/66008_reglement_20251030.pdf"
    hauteurs_url: str = "https://data.geopf.fr/annexes/gpu/documents/DU_66008/dd4c8deab39aa04c8938e88dd2337dba/66008_reglement_graphique_4_20251030.pdf"
    ppr_url: str = (
        "https://www.pyrenees-orientales.gouv.fr/contenu/telechargement/46345/357350/file/"
        "PM1_ArgelesSurMer_PPRn_20170529_reglement.pdf"
    )
    pprif_url: str = (
        "https://piece-jointe-carto.developpement-durable.gouv.fr/DEPT066A/DOC_PPRN/"
        "PM1_ArgelesSurMer_PPRif_20060627_reglement.pdf"
    )
    taxe_communale: str = "5 %"
    taxe_departementale: str = "2 %"
    rap: str = "0,40 %"

    maire: str = "Julie SANZ"
    delegation: str = "Didier WINZER, Responsable Service Urbanisme"

    mentions_communales: list = field(default_factory=lambda: [
        "La commune d'Argelès-sur-Mer est située dans une zone contaminée ou susceptible "
        "de l'être par les termites.",
        "La commune d'Argelès-sur-Mer est concernée par un risque d'exposition au plomb.",
        "La commune d'Argelès-sur-Mer est située en zone 3 de sismicité modérée.",
        "La commune d'Argelès-sur-Mer est située en zone 3 : zone à potentiel radon significatif.",
    ])


# Routage couche → section. Tout ce qui n'est pas listé tombe en "prescriptions".
# Couverture catalogue Argelès (catalogue_cua_argeles.json) :
#   dispositions → zonage_plu (via module zonage_plu), hauteurs (art. 3)
#   sup          → servitudes (art. 4, handler catalogue + module partagé)
#   risques      → ppr, pprif (via module ppr_et_pprif), alea_feu (via module alea_feu),
#                  retrait_gonflement_argiles_2026, old (art. 5)
#   prescriptions→ aoc, prescriptions_* (via module prescriptions_plu), infos_surf (hors DPU),
#                  haies_bocages, znieffs, zaer, batiments (art. 5/7)
#   métier       → reseaux_enedis_lineaires, prairies_et_natura_2000 (+ natura_2000,
#                  prairies_sensibles via module dédié), servitudes
LAYER_TO_SECTION = {
    # DPU : c'est une info de infos_surf filtrée, géré séparément (voir section_dpu)
    # SUP : réglementation via handler catalogue servitudes (voir section_sup)
    "ppr":                   "risques",
    "pprif":                 "risques",
    "retrait_gonflement_argiles_2026": "risques",
    "old":                   "risques",
    "alea_feu":              "risques",
    "zonage_plu":            "dispositions",
    "hauteurs":              "dispositions",
}
# Couches gérées par un rendu spécifique (pas dans le flux générique "objets")
LAYERS_METIER = {
    "reseaux_enedis_lineaires",
    "prairies_et_natura_2000",
    "natura_2000",
    "prairies_sensibles",
    "servitudes",
    "ppr",
    "pprif",
    "ppr_et_pprif",
    "alea_feu",
    "zonage_plu",
    "prescriptions_surf",
    "prescriptions_lineaires",
    "prescriptions_ponctuelles",
    "prescriptions_plu",
    "taxes",
}
# Couches prescriptions PLU gérées par prescriptions_plu (exclues du flux générique)
PRESCRIPTION_PLU_KEYS = frozenset({
    "prescriptions_surf",
    "prescriptions_lineaires",
    "prescriptions_ponctuelles",
})
# Statuts à ignorer silencieusement (ne rien montrer au pétitionnaire)
STATUTS_IGNORES = {"erreur", "table_absente"}
MIN_ZONAGE_PCT = 1.0

FONT = "Arial"
TITLE_BAR_FILL = "D9D9D9"


@dataclass
class Context:
    dossier: dict
    rapport: dict
    config: CommuneConfig


# ============================================================
# HELPERS DOCX
# ============================================================
def _set_cell_bg(cell, fill=TITLE_BAR_FILL):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_borders(cell, color="808080", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def add_title_bar(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell); _set_cell_borders(cell)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10); r.font.name = FONT
    doc.add_paragraph()


def add_para(doc, text="", bold=False, italic=False, size=10, align=None, space_after=4):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = FONT
    return p


def _add_logo_first_page(doc, logo_path: Path = LOGO_COMMUNE_PATH) -> None:
    """Logo commune en haut à droite de la première page."""
    if not logo_path.is_file():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        p.add_run().add_picture(str(logo_path), width=Cm(4.0))
    except Exception:
        pass


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        for c in cells: _set_cell_borders(c)
        cells[0].width = Cm(5.5); cells[1].width = Cm(11.5)
        r0 = cells[0].paragraphs[0].add_run(label); r0.bold = True
        r0.font.size = Pt(10); r0.font.name = FONT
        r1 = cells[1].paragraphs[0].add_run(str(value) if value not in (None, "") else "—")
        r1.font.size = Pt(10); r1.font.name = FONT
    doc.add_paragraph()


def _add_labeled_para(doc, label: str, value: str, *, size: int = 9, space_after: int = 3) -> None:
    """Ligne intitulée (label : valeur) pour les blocs PPR sous-zone."""
    text = str(value).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r_label = p.add_run(f"{label} : ")
    r_label.bold = True
    r_label.font.size = Pt(size)
    r_label.font.name = FONT
    r_val = p.add_run(text)
    r_val.font.size = Pt(size)
    r_val.font.name = FONT


def _add_parcelles_concernées_para(doc, parcelles: list[dict], *, size: int = 9) -> None:
    """Ligne « Parcelles concernées » en fin de sous-zone, références en gras."""
    if not parcelles:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r_label = p.add_run("Parcelles concernées : ")
    r_label.bold = True
    r_label.font.size = Pt(size)
    r_label.font.name = FONT

    first = True
    for parcelle in parcelles:
        ref = (parcelle.get("libelle") or "").strip()
        if not ref:
            section = (parcelle.get("section") or "").strip().upper()
            numero = (parcelle.get("numero") or "").strip()
            if section and numero:
                ref = f"Parcelle {section} n°{numero}"
        if not ref:
            continue

        if not first:
            r_sep = p.add_run(", ")
            r_sep.font.size = Pt(size)
            r_sep.font.name = FONT
        first = False

        r_ref = p.add_run(ref)
        r_ref.bold = True
        r_ref.font.size = Pt(size)
        r_ref.font.name = FONT

        try:
            pct = float(parcelle.get("pct") or 0)
        except (TypeError, ValueError):
            pct = 0.0
        if pct > 0:
            r_pct = p.add_run(f" ({pct:.2f} %)")
            r_pct.font.size = Pt(size)
            r_pct.font.name = FONT


def _write_ppr_sous_zone(doc, bloc: dict) -> None:
    """Paragraphes intitulés pour une sous-zone PPR (sans tableau)."""
    _add_labeled_para(doc, "Sous-zone", bloc.get("label") or "", size=10)
    _add_labeled_para(doc, "Risque", bloc.get("risque") or "")
    _add_labeled_para(doc, "Degré", bloc.get("degre") or "")
    _add_labeled_para(doc, "Coefficient d'emprise au sol (CES)", bloc.get("ces") or "")
    _add_labeled_para(doc, "Mise hors d'eau obligatoire", bloc.get("mise_hors_d_eau") or "")

    regl = (
        (bloc.get("reglementation_generale") or bloc.get("reglementation") or "")
        .strip()
    )
    if regl:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Règlementation de la sous-zone :")
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = FONT
        add_markdown_block(doc, regl, size=9)

    _add_parcelles_concernées_para(doc, bloc.get("parcelles") or [])

    doc.add_paragraph()


# ============================================================
# EXTRACTION DU TEXTE D'UN OBJET
# ============================================================
def _format_cadastre(dossier: dict, rapport: dict) -> Optional[str]:
    """Cadastre du dossier, ou liste des parcelles de l'UF depuis le rapport."""
    cadastre = dossier.get("cadastre")
    if cadastre and str(cadastre).strip():
        return str(cadastre).strip()
    parcelles = rapport.get("parcelles") or []
    if not parcelles:
        return None
    parts = []
    for p in parcelles:
        section = (p.get("section") or "").strip()
        numero = (p.get("numero") or "").strip()
        if section and numero:
            parts.append(f"{section} n°{numero}")
    return ", ".join(parts) if parts else None


def _format_adresses_parcelles(rapport: dict) -> Optional[str]:
    """Adresse(s) BAN liée(s) aux parcelles de l'UF (résolution locale)."""
    bloc = rapport.get("adresses_parcelles")
    if bloc is None:
        return None
    txt = (bloc.get("texte_header") or "").strip()
    return txt if txt else "—"


def _reglementation_text(obj: dict) -> Optional[str]:
    regl = obj.get("reglementation")
    if regl and str(regl).strip() and str(regl).strip() != "\\N":
        return str(regl).strip()
    return None


def texte_objet(obj: dict, *, skip_reglementation: bool = False) -> Optional[str]:
    """Réglementation taguée en priorité ; fallback selon les attributs métier de la couche."""
    if not skip_reglementation:
        regl = _reglementation_text(obj)
        if regl:
            return regl

    # Hauteurs PLU (libellé long / valeur numérique)
    if any(k in obj for k in ("hauteur", "libelong")):
        parts = []
        libelong = (obj.get("libelong") or "").strip()
        hauteur = obj.get("hauteur")
        if libelong:
            parts.append(libelong)
        if hauteur not in (None, ""):
            parts.append(f"hauteur maximale : {hauteur} m")
        if parts:
            return " — ".join(parts)

    fallback = (
        obj.get("libelle")
        or obj.get("libelong")
        or obj.get("legende")
        or obj.get("nomsuplitt")
        or obj.get("nom_site")
        or obj.get("denom")
        or obj.get("nom")
    )
    return str(fallback).strip() if fallback else None


def _is_dpu_objet(obj: dict) -> bool:
    lib = (obj.get("libelle") or "").lower()
    return "préemption" in lib or "preemption" in lib


def _cle_dedup_libelle(obj: dict) -> Optional[str]:
    """Clé de dédoublonnage : libellé métier, sinon texte réglementaire affiché."""
    lib = (obj.get("libelle") or "").strip()
    if lib:
        return lib.casefold()
    regl = _reglementation_text(obj)
    if regl:
        return regl.casefold()
    txt = texte_objet(obj, skip_reglementation=True)
    return txt.casefold() if txt else None


def _dedupe_objets_par_libelle(objets: list) -> list:
    """Conserve une seule entrée par libellé distinct (ordre d'apparition conservé)."""
    seen: set[str] = set()
    out: list = []
    for obj in objets:
        cle = _cle_dedup_libelle(obj)
        if cle is None:
            out.append(obj)
            continue
        if cle in seen:
            continue
        seen.add(cle)
        out.append(obj)
    return out


def _objets_affichables(key: str, layer: dict) -> list:
    """Filtre les objets d'une couche (ex. DPU déjà traité dans section_dpu)."""
    objets = layer.get("objets") or []
    if key == "infos_surf":
        objets = [o for o in objets if not _is_dpu_objet(o)]
        return _dedupe_objets_par_libelle(objets)
    return objets


def couches_par_section(rapport: dict):
    """Regroupe les couches concernées par section cible. Retourne {section: [(key, layer)]}."""
    groupes = {}
    for key, layer in rapport.get("intersections", {}).items():
        if key in LAYERS_METIER:
            continue
        if layer.get("status") in STATUTS_IGNORES:
            continue
        if not layer.get("objets"):
            continue
        section = LAYER_TO_SECTION.get(key, "prescriptions")
        groupes.setdefault(section, []).append((key, layer))
    return groupes


def _write_note_intersection_partielle(
    doc,
    obj: dict,
    *,
    multi_entites: bool,
    enabled: bool,
) -> None:
    note = note_pour_objet(obj, multi_entites=multi_entites, enabled=enabled)
    if note:
        add_para(doc, note, italic=True, size=9, space_after=3)


def _bloc_couche(doc, layer, prefix_nom=True, objets=None):
    """Écrit le nom de couche en gras + une puce par objet (texte réglementaire)."""
    rows = objets if objets is not None else layer.get("objets") or []
    if not rows:
        return
    enabled = bool(layer.get("afficher_pct_sig_partiel"))
    multi = est_multi_entites(layer.get("objets") or rows)
    if prefix_nom:
        add_para(doc, layer.get("nom") or "", bold=True, space_after=2)
    seen_regl: set[str] = set()
    for obj in rows:
        regl = _reglementation_text(obj)
        if regl:
            if regl in seen_regl:
                continue
            seen_regl.add(regl)
            add_markdown_block(doc, regl, size=9)
            _write_note_intersection_partielle(
                doc, obj, multi_entites=multi, enabled=enabled
            )
            continue
        txt = texte_objet(obj, skip_reglementation=True)
        if not txt:
            continue
        add_para(doc, "• " + txt, size=9, space_after=3)
        _write_note_intersection_partielle(
            doc, obj, multi_entites=multi, enabled=enabled
        )


def _pct_sig(obj: dict) -> float:
    try:
        return float(obj.get("pct_sig") or 0)
    except (TypeError, ValueError):
        return 0.0


def _label_obj(obj: dict, *keys: str) -> str:
    for key in keys:
        val = (obj.get(key) or "").strip()
        if val:
            return val
    return ""


def _items_avec_pct(objets: list, *label_keys: str) -> list[tuple[str, float]]:
    """Libellés distincts avec part UF maximale (> MIN_ZONAGE_PCT)."""
    seen: dict[str, float] = {}
    for obj in objets:
        pct = _pct_sig(obj)
        if pct <= MIN_ZONAGE_PCT:
            continue
        label = _label_obj(obj, *label_keys)
        if not label:
            continue
        seen[label] = max(seen.get(label, 0.0), pct)
    return sorted(seen.items(), key=lambda item: -item[1])


def _write_zonage_plu_detail_parcelles(doc, module: dict) -> bool:
    """Détail zonage par parcelle (UF multi-parcelles) — en tête de section PLU."""
    details = module.get("detail_parcelles") or []
    if not details:
        return False

    add_para(doc, "Détail par parcelle de l'unité foncière :", bold=True, space_after=4)
    for parcelle in details:
        texte = (parcelle.get("texte") or "").strip()
        if texte:
            add_para(doc, f"• {texte}", size=9, space_after=3)
    doc.add_paragraph()
    return True


def _write_zonage_plu_details(doc, module: dict) -> bool:
    """Zonage PLU : items pré-calculés par le module intersection_modules/zonage_plu."""
    items = module.get("items") or []
    if not items:
        return False

    for item in items:
        kind = item.get("kind")
        if kind == "reglementation":
            titre = (item.get("titre") or "").strip()
            if titre:
                add_para(doc, titre, bold=True, space_after=2)
            regl = (item.get("reglementation") or "").strip()
            if regl:
                add_markdown_block(doc, regl, size=9)
            continue
        if kind == "bullet":
            texte = (item.get("texte") or "").strip()
            if texte:
                add_para(doc, f"• {texte}", size=9, space_after=3)
    return True


def _format_hauteurs_intro(objets: list) -> tuple[list[str], Optional[str]]:
    """Résumé secteurs de hauteur + parts de surface significatives."""
    items = _items_avec_pct(objets, "libelong")
    if items:
        secteurs = [secteur for secteur, _ in items]
        if len(items) == 1:
            secteur, pct = items[0]
            texte = (
                f"L'unité foncière est soumise aux règles de hauteur du secteur {secteur} "
                f"({pct:.2f} % de la surface)."
            )
        else:
            parts = [f"{secteur} ({pct:.2f} %)" for secteur, pct in items]
            texte = (
                "L'unité foncière est soumise aux règles de hauteur des secteurs "
                f"{', '.join(parts)}."
            )
        return secteurs, texte

    secteurs, seen = [], set()
    for obj in objets:
        secteur = _label_obj(obj, "libelong")
        if secteur and secteur not in seen:
            seen.add(secteur)
            secteurs.append(secteur)
    if not secteurs:
        return [], None
    if len(secteurs) == 1:
        return secteurs, f"L'unité foncière est soumise aux règles de hauteur du secteur {secteurs[0]}."
    return secteurs, (
        "L'unité foncière est soumise aux règles de hauteur des secteurs "
        f"{', '.join(secteurs)}."
    )


def _objets_significatifs(objets: list) -> list:
    return [obj for obj in objets if _pct_sig(obj) > MIN_ZONAGE_PCT]


def _titre_hauteur_obj(obj: dict) -> Optional[str]:
    libelong = (obj.get("libelong") or "").strip()
    return libelong or None


def _write_hauteurs_details(
    doc,
    layer: dict,
    *,
    show_layer_title: bool = True,
    secteurs: Optional[list] = None,
):
    """Hauteurs PLU : libelong en titre + réglementation markdown."""
    objets = layer.get("objets") or []
    if not objets:
        return False

    if show_layer_title:
        add_para(doc, layer.get("nom") or "Hauteurs maximales (PLU)", bold=True, space_after=2)
    seen_regl: set[str] = set()
    multi_secteurs = len(secteurs or []) > 1
    wrote = False

    for obj in _objets_significatifs(objets):
        regl = _reglementation_text(obj)
        titre = _titre_hauteur_obj(obj)
        pct = _pct_sig(obj)

        if regl:
            if regl in seen_regl:
                continue
            seen_regl.add(regl)
            if titre:
                suffix = f" ({pct:.2f} %)" if multi_secteurs and pct > MIN_ZONAGE_PCT else ""
                add_para(doc, f"{titre}{suffix}", bold=True, size=9, space_after=2)
            add_markdown_block(doc, regl, size=9)
            wrote = True
            continue

        if titre:
            suffix = f" ({pct:.2f} %)" if multi_secteurs and pct > MIN_ZONAGE_PCT else ""
            add_para(doc, f"• {titre}{suffix}", size=9, space_after=3)
            wrote = True
            continue

        txt = texte_objet(obj, skip_reglementation=True)
        if txt:
            add_para(doc, f"• {txt}", size=9, space_after=3)
            wrote = True

    return wrote


# ============================================================
# SECTIONS
# ============================================================
def section_identite(doc, ctx):
    d = ctx.dossier
    superficie = d.get("superficie") or ctx.rapport.get("surface_indicative")
    rows = [
        ("Par", d.get("demandeur")),
        ("Demeurant à", d.get("demandeur_adresse")),
        ("Sur un terrain sis", d.get("terrain")),
        ("Cadastre", _format_cadastre(d, ctx.rapport)),
    ]
    adresses_parcelles = _format_adresses_parcelles(ctx.rapport)
    if "adresses_parcelles" in ctx.rapport:
        rows.append(("Adresse(s) de l'unité foncière", adresses_parcelles))
    rows.extend([
        ("Demande déposée le", d.get("date_depot") or datetime.now().strftime("%d/%m/%Y")),
        ("Superficie", f"{superficie} m²" if superficie else None),
        ("N° de dossier", d.get("numero_cu")),
    ])
    add_kv_table(doc, rows)


def section_carte_identite(doc, ctx):
    """Encart carte d'identité d'urbanisme (lien vers HTML gelé)."""
    url = ctx.dossier.get("carte_context_url") or ctx.rapport.get("carte_context_url")
    if not url:
        return
    add_title_bar(doc, "Carte d'identité d'urbanisme de l'unité foncière")
    add_para(
        doc,
        "Les informations géographiques et réglementaires applicables à l'unité foncière "
        "objet du présent certificat sont consultables sur la carte d'identité d'urbanisme "
        "associée, accessible à l'adresse suivante :",
        size=9,
    )
    add_para_with_link(doc, "", url, url, size=9)
    doc.add_paragraph()


def section_vu(doc, ctx):
    c = ctx.config
    add_title_bar(doc, "Demande en vue de connaître les dispositions d'urbanisme applicables au terrain")
    for ligne in [
        "Vu la demande de certificat d'urbanisme ;",
        "Vu le code de l'urbanisme ;",
        f"Vu le Plan local d'urbanisme {c.plu_mention} ;",
        f"Vu le Plan de Prévention des Risques Naturels Prévisibles {c.pprn_mention} ;",
        f"Vu le {c.pgri_mention}.",
    ]:
        add_para(doc, ligne)
    doc.add_paragraph()


def section_dpu(doc, ctx):
    add_title_bar(doc, "Droit de préemption")
    # Le DPU vient de infos_surf (ligne libelle='Droit de Préemption Urbain', taguée)
    dpu_txt = None
    for key in ("infos_surf",):
        layer = ctx.rapport.get("intersections", {}).get(key, {})
        for obj in layer.get("objets", []):
            lib = (obj.get("libelle") or "").lower()
            if "préemption" in lib or "preemption" in lib:
                dpu_txt = texte_objet(obj)
                break
    if dpu_txt:
        add_para(doc, dpu_txt)
    else:
        add_para(doc, "L'unité foncière n'est pas soumise au Droit de Préemption Urbain (DPU).")
    doc.add_paragraph()


def _format_ac1_monument_line(nom: str, dist: float | None) -> str:
    """Libellé AC1 : distance_m mesure l'écart au périmètre des abords (buffer), pas au monument."""
    if dist is None:
        return nom
    if dist <= 0:
        return f"{nom} — l'unité foncière est concernée par le périmètre des abords"
    return f"{nom} — à environ {dist:.0f} m du périmètre des abords"


def _write_ac1_monuments(doc, servitude: dict) -> None:
    """Monuments historiques AC1 : nom en gras + distance au périmètre des abords."""
    monuments = servitude.get("monuments")
    if monuments:
        for mon in monuments:
            nom = (mon.get("nom") or "").strip()
            if not nom:
                continue
            add_para(
                doc,
                _format_ac1_monument_line(nom, mon.get("distance_m")),
                bold=True,
                size=9,
                space_after=2,
            )
        return

    if str(servitude.get("suptype") or "").strip().lower() != "ac1":
        return
    nom = (servitude.get("nomsuplitt") or "").strip()
    if not nom:
        return
    add_para(
        doc,
        _format_ac1_monument_line(nom, servitude.get("distance_m")),
        bold=True,
        size=9,
        space_after=2,
    )


def _dedupe_servitudes(servitudes: list[dict]) -> list[dict]:
    """Une entrée par suptype (agrégation déjà faite dans modules_communs)."""
    seen: set[str] = set()
    out: list[dict] = []
    for s in servitudes:
        key = (s.get("suptype") or "").strip().upper()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(s)
    return out


def _write_i4_variantes(doc, servitude: dict) -> None:
    variantes = servitude.get("variantes") or []
    if not variantes:
        return
    base_regl = (servitude.get("reglementation") or "").strip()
    if base_regl:
        add_markdown_block(doc, base_regl, size=9)
    for var in variantes:
        libelle_var = (var.get("libelle_var") or "").strip()
        if libelle_var:
            add_para(doc, libelle_var, bold=True, size=9, space_after=2)
        complement = (var.get("complement") or "").strip()
        if complement:
            add_markdown_block(doc, complement, size=9)
        var_metric = var.get("metric")
        var_nb = var.get("nb_fragments")
        if var_metric is not None and var_nb and var_nb > 1:
            add_para(
                doc,
                f"Surface d'intersection : {var_metric:,.2f} m² ({var_nb} fragment(s)).",
                italic=True,
                size=9,
                space_after=2,
            )
        elif var_metric is not None:
            add_para(
                doc,
                f"Surface d'intersection : {var_metric:,.2f} m².",
                italic=True,
                size=9,
                space_after=2,
            )


def section_sup(doc, ctx):
    inter = ctx.rapport.get("intersections", {})
    serv_layer = inter.get("servitudes") or inter.get("servitudes_reglementees", {})
    servitudes = _dedupe_servitudes(serv_layer.get("servitudes") or [])
    if not servitudes:
        return
    enabled_partiel = bool(serv_layer.get("afficher_pct_sig_partiel"))
    multi_couche = len(servitudes) > 1
    surface_sig = float(ctx.rapport.get("surface_m2") or 0)
    add_title_bar(doc, "Servitudes d'utilité publique")
    for i, s in enumerate(servitudes):
        if i > 0:
            doc.add_paragraph()
        titre = s.get("libelle") or s.get("nomsuplitt") or s.get("suptype") or "Servitude"
        add_para(doc, titre, bold=True, space_after=4)
        _write_ac1_monuments(doc, s)
        variantes = s.get("variantes") or []
        if variantes:
            _write_i4_variantes(doc, s)
        else:
            regl = (s.get("reglementation") or "").strip()
            if regl:
                add_markdown_block(doc, regl, size=9)
        note = note_pour_servitude(
            s,
            multi_entites_couche=multi_couche,
            enabled=enabled_partiel,
            surface_sig=surface_sig,
        )
        if note:
            add_para(doc, note, italic=True, size=9, space_after=4)
        base = (s.get("base_legale") or "").strip()
        if base:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            parse_markdown_inline(p, f"Base légale : {base}", size=9, italic_base=True)
        url_gpu = (s.get("url_fiche_gpu") or "").strip()
        if url_gpu:
            add_para_with_link(
                doc,
                "Fiche GPU : ",
                url_gpu,
                url_gpu,
                size=9,
                space_after=4,
            )
    doc.add_paragraph()


def _write_ppr_pprif_detail_parcelles(doc, module: dict) -> bool:
    """Détail PPR / PPRIF par parcelle (UF multi-parcelles) — si pas déjà sur les blocs."""
    ppr_blocs = (module.get("ppr") or {}).get("blocs") or []
    pprif_blocs = (module.get("pprif") or {}).get("blocs") or []
    if any(b.get("parcelles") for b in ppr_blocs + pprif_blocs):
        return False

    details = module.get("detail_parcelles") or []
    if not details:
        return False

    add_para(doc, "Détail par parcelle de l'unité foncière :", bold=True, space_after=4)
    for parcelle in details:
        texte = (parcelle.get("texte") or "").strip()
        if texte:
            add_para(doc, f"• {texte}", size=9, space_after=3)
    doc.add_paragraph()
    return True


def _write_ppr_pprif_details(doc, module: dict, config: CommuneConfig) -> bool:
    """PPR / PPRIF : paragraphes intitulés par sous-zone (PPR) + laius PPRIF."""
    ppr_data = module.get("ppr") or {}
    pprif_data = module.get("pprif") or {}
    ppr_blocs = ppr_data.get("blocs") or []
    pprif_blocs = pprif_data.get("blocs") or []
    has_detail = bool(module.get("detail_parcelles"))
    if not ppr_blocs and not pprif_blocs and not has_detail:
        return False

    _write_ppr_pprif_detail_parcelles(doc, module)

    if ppr_blocs:
        add_para(doc, ppr_data.get("nom") or "PPR (Plan de Prévention des Risques)", bold=True, space_after=2)
        add_para(
            doc,
            "L'unité foncière est concernée par une ou plusieurs zones du Plan de Prévention des Risques :",
            size=9,
            space_after=6,
        )
        for bloc in ppr_blocs:
            _write_ppr_sous_zone(doc, bloc)
        seen_note_codes: set[str] = set()
        for note in ppr_data.get("notes") or []:
            code = (note.get("code") or "").strip().upper()
            if code and code in seen_note_codes:
                continue
            note_regl = (note.get("reglementation") or "").strip()
            if not note_regl:
                continue
            if code:
                seen_note_codes.add(code)
            add_markdown_block(doc, note_regl, size=9)
        add_para_with_link(
            doc,
            "Règlement PPRN consultable sur : ",
            config.ppr_url,
            config.ppr_url,
            space_after=6,
        )

    if pprif_blocs:
        add_para(doc, pprif_data.get("nom") or "PPRIF (Risque Incendie de Forêt)", bold=True, space_after=2)
        for bloc in pprif_blocs:
            rows = []
            if bloc.get("risque"):
                rows.append(("Risque", bloc["risque"]))
            if bloc.get("zone"):
                rows.append(("Zone", bloc["zone"]))
            if rows:
                add_kv_table(doc, rows)
            regl = (bloc.get("reglementation") or "").strip()
            if regl:
                add_markdown_block(doc, regl, size=9)
            _add_parcelles_concernées_para(doc, bloc.get("parcelles") or [])
        add_para_with_link(
            doc,
            "Règlement PPRIF consultable sur : ",
            config.pprif_url,
            config.pprif_url,
            space_after=6,
        )

    return True


def _write_alea_feu_details(doc, module: dict) -> bool:
    """PAC — aléas incendie de forêt et de végétation par libellé distinct."""
    blocs = module.get("blocs") or []
    if not blocs:
        return False

    add_para(
        doc,
        module.get("nom") or "Risque d'incendie de forêt et de végétation",
        bold=True,
        space_after=2,
    )
    intro = (module.get("intro") or "").strip()
    if intro:
        add_para(doc, intro, size=9, space_after=6)
    for bloc in blocs:
        libelle = (bloc.get("libelle") or "").strip()
        if not libelle:
            continue
        pct = float(bloc.get("pct_sig") or 0)
        add_para(
            doc,
            f"• {libelle} ({pct:.2f} % de l'unité foncière)",
            size=9,
            space_after=3,
        )
    pac_url = (module.get("pac_url") or "").strip()
    if pac_url:
        add_para_with_link(
            doc,
            "Porté à connaissance consultable sur : ",
            pac_url,
            pac_url,
            space_after=6,
        )
    return True


def section_risques(doc, ctx):
    groupes = couches_par_section(ctx.rapport)
    layers = groupes.get("risques", [])
    pp = ctx.rapport.get("intersections", {}).get("ppr_et_pprif", {})
    alea_feu = ctx.rapport.get("intersections", {}).get("alea_feu", {})
    has_pp = bool(
        (pp.get("ppr") or {}).get("blocs")
        or (pp.get("pprif") or {}).get("blocs")
        or pp.get("detail_parcelles")
    )
    has_alea_feu = bool(alea_feu.get("blocs"))
    if not layers and not has_pp and not has_alea_feu:
        return
    add_title_bar(doc, "Risques naturels et technologiques")
    if has_pp:
        _write_ppr_pprif_details(doc, pp, ctx.config)
    if has_alea_feu:
        _write_alea_feu_details(doc, alea_feu)
    for key, layer in layers:
        _bloc_couche(doc, layer)
    doc.add_paragraph()


def _write_prescriptions_plu_detail_parcelles(doc, module: dict) -> bool:
    """Détail prescriptions PLU par parcelle (libellés touchés, sans %)."""
    details = module.get("detail_parcelles") or []
    if not details:
        return False

    add_para(doc, "Détail par parcelle de l'unité foncière :", bold=True, space_after=4)
    for parcelle in details:
        texte = (parcelle.get("texte") or "").strip()
        if texte:
            add_para(doc, f"• {texte}", size=9, space_after=3)
    doc.add_paragraph()
    return True


def _write_prescriptions_plu_item_note(doc, item: dict, *, multi_entites: bool, enabled: bool) -> None:
    if not enabled:
        return
    try:
        pct = float(item.get("pct_sig") or 0)
    except (TypeError, ValueError):
        return
    note = texte_note_intersection_partielle(
        pct, multi_entites=multi_entites, enabled=True
    )
    if note:
        add_para(doc, note, italic=True, size=9, space_after=3)


def _write_prescriptions_plu_details(doc, module: dict) -> bool:
    """Prescriptions PLU : libellé + réglementation ; note partielle si catalogue activé."""
    couches = module.get("couches") or []
    has_detail = bool(module.get("detail_parcelles"))
    if not couches and not has_detail:
        return False

    _write_prescriptions_plu_detail_parcelles(doc, module)

    for couche in couches:
        nom = (couche.get("nom") or "").strip()
        items = couche.get("items") or []
        if not items:
            continue
        enabled_partiel = bool(couche.get("afficher_pct_sig_partiel"))
        multi = (couche.get("nb_objets") or 0) > 1
        if nom:
            add_para(doc, nom, bold=True, space_after=2)
        for item in items:
            kind = item.get("kind")
            if kind == "reglementation":
                libelle = (item.get("libelle") or "").strip()
                if libelle:
                    add_para(doc, libelle, bold=True, size=9, space_after=2)
                regl = (item.get("reglementation") or "").strip()
                if regl:
                    add_markdown_block(doc, regl, size=9)
                _write_prescriptions_plu_item_note(
                    doc, item, multi_entites=multi, enabled=enabled_partiel
                )
                continue
            if kind == "bullet":
                texte = (item.get("texte") or "").strip()
                if texte:
                    add_para(doc, f"• {texte}", size=9, space_after=3)
                _write_prescriptions_plu_item_note(
                    doc, item, multi_entites=multi, enabled=enabled_partiel
                )
    return True


def section_prescriptions(doc, ctx):
    prescriptions_plu = ctx.rapport.get("intersections", {}).get("prescriptions_plu", {})
    groupes = couches_par_section(ctx.rapport)
    layers = [
        (key, layer)
        for key, layer in groupes.get("prescriptions", [])
        if key not in PRESCRIPTION_PLU_KEYS
    ]
    has_prescriptions_plu = bool(
        prescriptions_plu.get("couches") or prescriptions_plu.get("detail_parcelles")
    )
    if not layers and not has_prescriptions_plu:
        return
    add_title_bar(doc, "Prescriptions et informations applicables au terrain")
    if has_prescriptions_plu:
        _write_prescriptions_plu_details(doc, prescriptions_plu)
    for key, layer in layers:
        _bloc_couche(doc, layer, objets=_objets_affichables(key, layer))
    doc.add_paragraph()


def _write_natura_partielles(doc, natura_layer: dict) -> None:
    """Notes d'intersection partielle Natura 2000 (couche catalogue, section métier)."""
    if not natura_layer or not natura_layer.get("afficher_pct_sig_partiel"):
        return
    objets = natura_layer.get("objets") or []
    if not objets:
        return
    multi = est_multi_entites(objets)
    for obj in objets:
        note = note_pour_objet(obj, multi_entites=multi, enabled=True)
        if not note:
            continue
        site = _label_obj(obj, "n_site", "c_site", "id")
        if site:
            add_para(doc, f"• {site} — {note}", italic=True, size=9, space_after=3)
        else:
            add_para(doc, note, italic=True, size=9, space_after=3)


def _write_prairies_natura_details(doc, pn: dict):
    """Natura 2000 / Prairies : blocs réglementaires depuis la table dédiée."""
    if not (pn.get("has_natura") or pn.get("has_prairie")):
        return

    add_title_bar(doc, "Natura 2000 / Prairies sensibles")
    diag = (pn.get("diagnostic_metier") or "").strip()
    if diag:
        add_para(doc, diag, bold=True, space_after=4)

    blocs = pn.get("blocs") or []
    if not blocs:
        for legacy in (pn.get("natura"), pn.get("prairie")):
            if not legacy:
                continue
            regl = (legacy.get("reglementation") or legacy.get("laius") or "").strip()
            if regl:
                blocs.append(legacy)

    for bloc in blocs:
        titre = (bloc.get("nom_regime") or bloc.get("code_regime") or "").strip()
        if titre:
            add_para(doc, titre, bold=True, space_after=2)
        statut = (bloc.get("statut_juridique") or bloc.get("statut") or "").strip()
        if statut:
            add_para(doc, f"Statut juridique : {statut}", size=9, italic=True, space_after=2)
        regl = (bloc.get("reglementation") or bloc.get("laius") or "").strip()
        if regl:
            add_markdown_block(doc, regl, size=9)
        base = (bloc.get("base_legale") or "").strip()
        if base:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            parse_markdown_inline(p, f"Base légale : {base}", size=9, italic_base=True)

    doc.add_paragraph()


def section_metier(doc, ctx):
    """Modules métier (ENEDIS, Natura/Prairies) : rendu via leurs champs dédiés."""
    inter = ctx.rapport.get("intersections", {})

    enedis = inter.get("reseaux_enedis_lineaires", {})
    analyses = enedis.get("analyses") or []
    if analyses:
        add_title_bar(doc, "Réseaux électriques BT à proximité (ENEDIS)")
        add_para(
            doc,
            "Indications issues des données linéaires ENEDIS et d'analyses SIG "
            "(distance au câble le plus proche). Ne constitue pas une étude de "
            "faisabilité de raccordement.",
            size=9,
            italic=True,
            space_after=6,
        )
        for a in analyses:
            diag = a.get("diagnostic_expert_raccordement")
            if diag:
                add_para(doc, f"• {diag}", size=9, space_after=3)
        doc.add_paragraph()

    pn = inter.get("prairies_et_natura_2000", {})
    _write_prairies_natura_details(doc, pn)
    _write_natura_partielles(doc, inter.get("natura_2000", {}))


def section_dispositions(doc, ctx):
    c = ctx.config
    add_title_bar(doc, "Nature des dispositions d'urbanisme applicables au terrain")
    for ligne in [
        "Articles d'ordre public du Règlement National d'Urbanisme : R.111-2, R.111-4, "
        "R.111-25 à R.111-27 et R.111-31 à R.111-51 du Code de l'urbanisme.",
        "Articles L.111-6 à L.111-10 du Code de l'urbanisme.",
        "Loi Littoral n° 86-2 du 3 janvier 1986.",
        "Loi Montagne n° 85-30 du 9 janvier 1985.",
        f"Plan local d'urbanisme (PLU) {c.plu_mention}.",
    ]:
        add_para(doc, ligne)

    doc.add_paragraph()

    # ── Zonage PLU ──
    zonage = ctx.rapport.get("intersections", {}).get("zonage_plu", {})
    intro_zonage = zonage.get("intro")

    add_title_bar(doc, "Zonage du PLU")
    _write_zonage_plu_detail_parcelles(doc, zonage)
    if intro_zonage:
        add_para(doc, intro_zonage, bold=True, space_after=6)
    else:
        add_para(doc, "Zonage PLU non déterminé pour ce terrain.", italic=True, space_after=6)
    _write_zonage_plu_details(doc, zonage)
    add_para_with_link(
        doc,
        "Règlement PLU consultable sur : ",
        c.geoportail_url,
        c.geoportail_url,
        space_after=6,
    )

    # ── Hauteurs PLU ──
    groupes = couches_par_section(ctx.rapport)
    layer_hauteurs = next(
        (layer for key, layer in groupes.get("dispositions", []) if key == "hauteurs"),
        None,
    )
    if layer_hauteurs and (layer_hauteurs.get("objets") or []):
        objets_hauteurs = layer_hauteurs.get("objets") or []
        secteurs, intro_hauteurs = _format_hauteurs_intro(objets_hauteurs)
        doc.add_paragraph()
        add_title_bar(doc, "Réglementation liée aux hauteurs")
        if intro_hauteurs:
            add_para(doc, intro_hauteurs, bold=True, space_after=6)
        _write_hauteurs_details(
            doc,
            layer_hauteurs,
            show_layer_title=False,
            secteurs=secteurs,
        )
        add_para_with_link(
            doc,
            "Carte des hauteurs consultable sur : ",
            c.hauteurs_url,
            c.hauteurs_url,
            space_after=6,
        )

    # Autres couches art. 3 éventuelles
    for key, layer in groupes.get("dispositions", []):
        if key in ("zonage_plu", "hauteurs"):
            continue
        _bloc_couche(doc, layer, objets=_objets_affichables(key, layer))

    doc.add_paragraph()


def section_sursis(doc, ctx):
    add_title_bar(doc, "Sursis à statuer")
    add_para(doc, "La commune peut décider de surseoir à statuer, dans les conditions et délais "
                  "prévus à l'article L.424-1 du code de l'urbanisme, sur les demandes d'autorisation "
                  "qui seraient de nature à compromettre ou rendre plus onéreuse l'exécution du futur PLU.")
    doc.add_paragraph()


def section_taxes(doc, ctx):
    c = ctx.config
    taxes = ctx.rapport.get("intersections", {}).get("taxes", {})
    taxe_communale = taxes.get("taux_communale_libelle") or c.taxe_communale
    add_title_bar(doc, "Taxes et contributions")
    add_para(doc, "Les taxes et contributions ne peuvent être déterminées précisément qu'à l'examen "
                  "de la demande d'autorisation.")
    add_para(doc, "Fiscalité applicable au terrain :", bold=True)
    add_para(doc, f"– Taxe d'aménagement – part communale (taux : {taxe_communale})")
    add_para(doc, f"– Taxe d'aménagement – part départementale (taux : {c.taxe_departementale})")
    add_para(doc, f"– Redevance d'archéologie préventive (taux : {c.rap})")
    add_para(doc, "Participations applicables au terrain :", bold=True)
    add_para(doc, "– Projet Urbain Partenarial")
    add_para(doc, "– Participation pour équipements publics exceptionnels")
    doc.add_paragraph()


def section_formalites(doc, ctx):
    add_title_bar(doc, "Formalités administratives préalables à l'opération")
    add_para(doc, "Préalablement à l'édification de construction ou à la réalisation de l'opération "
                  "projetée, les formalités ci-après devront être accomplies : Permis de Construire, "
                  "Permis d'Aménager, Déclaration Préalable, Permis de Démolir.")
    add_para(doc, "Attention : le non-respect de ces formalités ou l'utilisation du sol en "
                  "méconnaissance des règles indiquées est passible de l'amende prévue à "
                  "l'article L.480-4 du Code de l'urbanisme.", italic=True)
    doc.add_paragraph()


def section_signature(doc, ctx):
    c = ctx.config
    add_para(doc, f"{c.nom.title()}, le {datetime.now().strftime('%d/%m/%Y')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, f"Le Maire, {c.maire}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Pour le Maire, par délégation", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, c.delegation, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    add_para(doc, "Le présent certificat est transmis au représentant de l'État (article L.2131-1 du CGCT).",
             size=9, italic=True)


def section_mentions(doc, ctx):
    for m in ctx.config.mentions_communales:
        add_para(doc, f"– {m}", size=9)
    doc.add_paragraph()


def section_informations(doc, ctx):
    add_title_bar(doc, "Informations")
    add_para(doc, "Durée de validité :", bold=True, size=9)
    add_para(doc, "Le certificat est valable dix-huit mois. Une demande déposée dans ce délai bénéficie "
                  "de la cristallisation des règles d'urbanisme, taxes et participations en vigueur à la "
                  "date du certificat, sauf dispositions de sécurité ou de salubrité publique.", size=9)
    add_para(doc, "Prolongation (article R.410-17 du code de l'urbanisme) :", bold=True, size=9)
    add_para(doc, "Prorogeable par période d'un an, sur demande deux mois avant l'expiration, si les règles "
                  "applicables n'ont pas évolué.", size=9)
    add_para(doc, "Délais et voies de recours :", bold=True, size=9)
    add_para(doc, "Recours contentieux possible devant le tribunal administratif dans les deux mois suivant "
                  "la notification (www.telerecours.fr).", size=9)


# Ordre du document
SECTIONS = [
    section_identite,
    section_carte_identite,
    section_vu,
    section_dpu,
    section_dispositions,
    section_sup,
    section_risques,
    section_prescriptions,
    section_metier,
    section_sursis,
    section_taxes,
    section_formalites,
    section_signature,
    section_mentions,
    section_informations,
]


# ============================================================
# POINT D'ENTRÉE
# ============================================================
def _setup_document() -> Document:
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10)
    s = doc.sections[0]
    s.page_height = Cm(29.7); s.page_width = Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Cm(2.0))
    return doc


def build_cua(dossier: dict, rapport: dict, output_path: str,
              config: Optional[CommuneConfig] = None) -> str:
    config = config or CommuneConfig()
    ctx = Context(dossier=dossier, rapport=rapport, config=config)
    doc = _setup_document()
    _add_logo_first_page(doc)
    add_para(doc, "CERTIFICAT D'URBANISME", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "délivré par le Maire au nom de la commune", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for section in SECTIONS:
        section(doc, ctx)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse, json
    ap = argparse.ArgumentParser(description="Builder CUA Argelès (v2 — consomme reglementation)")
    ap.add_argument("--rapport", required=True)
    ap.add_argument("--dossier", default=None)
    ap.add_argument("--output", default="CUA_argeles.docx")
    args = ap.parse_args()
    rapport = json.loads(open(args.rapport, encoding="utf-8").read())
    dossier = json.loads(open(args.dossier, encoding="utf-8").read()) if args.dossier else {}
    print("✅ CUA généré :", build_cua(dossier, rapport, args.output))