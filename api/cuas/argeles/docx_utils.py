# -*- coding: utf-8 -*-
"""docx_utils.py — Rendu Markdown → python-docx (gras, italique, listes, titres)."""

from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FONT = "Arial"
_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")

_NATIVE_BULLET_PREFIXES = ("•", "–", "—", "·")


def _strip_links(text: str) -> str:
    """Remplace [texte](url) par 'texte (url)' en texte brut."""
    return _LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)


def _style_run(run, size: int = 9, italic_base: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    if italic_base:
        run.italic = True


def parse_markdown_inline(paragraph, text, size=9, italic_base=False):
    """Parse **gras** et *italique* inline en runs Word (scan gauche → droite)."""
    text = _strip_links(text)
    pos = 0
    while pos < len(text):
        rest = text[pos:]
        bold_m = _BOLD_RE.search(rest)
        italic_m = _ITALIC_RE.search(rest)

        match = None
        if bold_m and (italic_m is None or bold_m.start() <= italic_m.start()):
            match = ("bold", bold_m)
        elif italic_m:
            match = ("italic", italic_m)

        if match is None:
            if rest:
                run = paragraph.add_run(rest)
                _style_run(run, size, italic_base)
            break

        kind, m = match
        if m.start() > 0:
            run = paragraph.add_run(rest[: m.start()])
            _style_run(run, size, italic_base)
        run = paragraph.add_run(m.group(1))
        _style_run(run, size, italic_base)
        if kind == "bold":
            run.bold = True
        else:
            run.italic = True
        pos += m.end()


def add_markdown_block(doc, text, size=9, is_bullet=False):
    """Ajoute un bloc multi-lignes Markdown dans le document.

    is_bullet=True force une puce Word sur chaque ligne — à éviter si le texte
    contient déjà des puces (•, –, etc.) ou des préfixes markdown (- / *).
    """
    if not text or not str(text).strip():
        return

    for raw_line in str(text).strip().split("\n"):
        line = raw_line.strip()
        if not line or line == "---":
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(max(size + 3, 11))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(max(size + 4, 12))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(max(size + 5, 13))
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            continue

        has_native_bullet = line.startswith(_NATIVE_BULLET_PREFIXES)
        is_line_bullet = is_bullet and not has_native_bullet
        cleaned_line = line
        if cleaned_line.startswith("- ") or cleaned_line.startswith("* "):
            is_line_bullet = True
            cleaned_line = cleaned_line[2:].strip()

        if is_line_bullet:
            p = doc.add_paragraph(style="List Bullet")
        else:
            p = doc.add_paragraph()

        p.paragraph_format.space_after = Pt(3)
        parse_markdown_inline(p, cleaned_line, size=size)


def add_hyperlink(paragraph, text: str, url: str, *, size: int = 10, italic: bool = False):
    """Insère un hyperlien cliquable dans un paragraphe existant."""
    part = paragraph.part
    r_id = part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if italic:
        italic_el = OxmlElement("w:i")
        r_pr.append(italic_el)

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), FONT)
    font.set(qn("w:hAnsi"), FONT)
    r_pr.append(font)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para_with_link(
    doc,
    prefix: str,
    link_text: str,
    url: str,
    *,
    size: int = 10,
    space_after: int = 4,
):
    """Paragraphe avec libellé + hyperlien cliquable."""
    p = doc.add_paragraph()
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if prefix:
        run = p.add_run(prefix)
        run.font.name = FONT
        run.font.size = Pt(size)
    add_hyperlink(p, link_text, url, size=size)
    return p
