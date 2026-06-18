# -*- coding: utf-8 -*-
"""
builder.py — Génération du Certificat d'Urbanisme (DOCX) pour Argelès-sur-Mer.

Consomme le rapport d'intersections (sortie de intersections.py) : pour chaque
couche concernée, lit la réglementation déjà taguée en base (obj["reglementation"])
et l'écrit dans la bonne section du document. Aucune logique réglementaire ici —
le texte vit en base, le builder ne fait que router et mettre en forme.

Architecture : une fonction par section, le builder itère sur SECTIONS.
- Changer de commune  → nouveau CommuneConfig.
- L'en-tête (logo / n° dossier / pagination) sera branchée séparément.
"""

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# CONFIG COMMUNE
# ============================================================
@dataclass
class CommuneConfig:
    nom: str = "ARGELÈS-SUR-MER"
    code_insee: str = "66008"
    departement: str = "Pyrénées-Orientales"

    plu_mention: str = ("approuvé le 20/04/2017, révisé le 10/03/2022, "
                        "modifié le 14/12/2023 et le 30/10/2025")
    pprn_mention: str = ("approuvé par arrêté préfectoral du 25/11/2008, "
                         "modifié par arrêté préfectoral du 29/05/2017")
    pgri_mention: str = ("porter à connaissance relatif aux règles de gestion du risque "
                         "d'inondation (PGRI) en date du 11/07/2019")
    geoportail_url: str = "https://www.geoportail-urbanisme.gouv.fr"

    taxe_communale: str = "5 %"
    taxe_departementale: str = "2 %"
    rap: str = "0,40 %"

    maire: str = "Julie SANZ"
    delegation: str = "Didier WINZER, Responsable Service Urbanisme"

    mentions_communales: list = field(default_factory=lambda: [
        "La commune d'Argelès-sur-Mer est située dans une zone contaminée ou susceptible "
        "de l'être par les termites.",
        "La commune d'Argelès-sur-Mer est concernée par un risque d'exposition au plomb.",
        "La commune d'Argelès-sur-Mer est située en zone 3 de sismicité modérée.",
        "La commune d'Argelès-sur-Mer est située en zone 3 : zone à potentiel radon significatif.",
    ])


# Routage couche → section. Tout ce qui n'est pas listé tombe en "prescriptions".
# Couverture catalogue Argelès (catalogue_cua_argeles.json) :
#   dispositions → zonage_plu, hauteurs (art. 3)
#   sup          → sup_* , generateurs_sup_lineaires (art. 4)
#   risques      → ppr, pprif, retrait_gonflement_argiles_2026, old (art. 5)
#   prescriptions→ aoc, prescriptions_*, infos_surf (hors DPU), haies_bocages,
#                  znieffs, prairies_sensibles, natura_2000, zaer, batiments (art. 5/7)
#   métier       → reseaux_enedis_lineaires, prairies_et_natura_2000 (dédié)
LAYER_TO_SECTION = {
    # DPU : c'est une info de infos_surf filtrée, géré séparément (voir section_dpu)
    "sup_assiette_s":        "sup",
    "sup_assiette_l":        "sup",
    "sup_assiette_p":        "sup",
    "sup_generateur_s":      "sup",
    "sup_generateur_l":      "sup",
    "sup_generateur_p":      "sup",
    "generateurs_sup_lineaires": "sup",
    "ppr":                   "risques",
    "pprif":                 "risques",
    "retrait_gonflement_argiles_2026": "risques",
    "old":                   "risques",
    "zonage_plu":            "dispositions",
    "hauteurs":              "dispositions",
}
# Couches gérées par un rendu spécifique (pas dans le flux générique "objets")
LAYERS_METIER = {"reseaux_enedis_lineaires", "prairies_et_natura_2000"}
# Statuts à ignorer silencieusement (ne rien montrer au pétitionnaire)
STATUTS_IGNORES = {"erreur", "table_absente"}

FONT = "Arial"
TITLE_BAR_FILL = "D9D9D9"


@dataclass
class Context:
    dossier: dict
    rapport: dict
    config: CommuneConfig


# ============================================================
# HELPERS DOCX
# ============================================================
def _set_cell_bg(cell, fill=TITLE_BAR_FILL):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_borders(cell, color="808080", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def add_title_bar(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell); _set_cell_borders(cell)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10); r.font.name = FONT
    doc.add_paragraph()


def add_para(doc, text="", bold=False, italic=False, size=10, align=None, space_after=4):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = FONT
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        for c in cells: _set_cell_borders(c)
        cells[0].width = Cm(5.5); cells[1].width = Cm(11.5)
        r0 = cells[0].paragraphs[0].add_run(label); r0.bold = True
        r0.font.size = Pt(10); r0.font.name = FONT
        r1 = cells[1].paragraphs[0].add_run(str(value) if value not in (None, "") else "—")
        r1.font.size = Pt(10); r1.font.name = FONT
    doc.add_paragraph()


# ============================================================
# EXTRACTION DU TEXTE D'UN OBJET
# ============================================================
def _format_cadastre(dossier: dict, rapport: dict) -> Optional[str]:
    """Cadastre du dossier, ou liste des parcelles de l'UF depuis le rapport."""
    cadastre = dossier.get("cadastre")
    if cadastre and str(cadastre).strip():
        return str(cadastre).strip()
    parcelles = rapport.get("parcelles") or []
    if not parcelles:
        return None
    parts = []
    for p in parcelles:
        section = (p.get("section") or "").strip()
        numero = (p.get("numero") or "").strip()
        if section and numero:
            parts.append(f"{section} n°{numero}")
    return ", ".join(parts) if parts else None


def texte_objet(obj: dict) -> Optional[str]:
    """Réglementation taguée en priorité ; fallback selon les attributs métier de la couche."""
    regl = obj.get("reglementation")
    if regl and str(regl).strip() and str(regl).strip() != "\\N":
        return str(regl).strip()

    # Hauteurs PLU (legende / libellé long / valeur)
    if any(k in obj for k in ("legende", "hauteur", "libelong")):
        parts = []
        legende = (obj.get("legende") or "").strip()
        libelong = (obj.get("libelong") or "").strip()
        hauteur = obj.get("hauteur")
        if legende:
            parts.append(legende)
        if hauteur not in (None, ""):
            parts.append(f"hauteur maximale : {hauteur} m")
        if libelong and libelong not in parts and libelong not in legende:
            parts.append(libelong)
        if parts:
            return " — ".join(parts)

    fallback = (
        obj.get("libelle")
        or obj.get("libelong")
        or obj.get("legende")
        or obj.get("nomsuplitt")
        or obj.get("nom_site")
        or obj.get("denom")
        or obj.get("nom")
    )
    return str(fallback).strip() if fallback else None


def _is_dpu_objet(obj: dict) -> bool:
    lib = (obj.get("libelle") or "").lower()
    return "préemption" in lib or "preemption" in lib


def _objets_affichables(key: str, layer: dict) -> list:
    """Filtre les objets d'une couche (ex. DPU déjà traité dans section_dpu)."""
    objets = layer.get("objets") or []
    if key == "infos_surf":
        return [o for o in objets if not _is_dpu_objet(o)]
    return objets


def couches_par_section(rapport: dict):
    """Regroupe les couches concernées par section cible. Retourne {section: [(key, layer)]}."""
    groupes = {}
    for key, layer in rapport.get("intersections", {}).items():
        if key in LAYERS_METIER:
            continue
        if layer.get("status") in STATUTS_IGNORES:
            continue
        if not layer.get("objets"):
            continue
        section = LAYER_TO_SECTION.get(key, "prescriptions")
        groupes.setdefault(section, []).append((key, layer))
    return groupes


def _bloc_couche(doc, layer, prefix_nom=True, objets=None):
    """Écrit le nom de couche en gras + une puce par objet (texte réglementaire)."""
    rows = objets if objets is not None else layer.get("objets") or []
    if not rows:
        return
    if prefix_nom:
        add_para(doc, layer.get("nom") or "", bold=True, space_after=2)
    for obj in rows:
        txt = texte_objet(obj)
        if not txt:
            continue
        add_para(doc, "• " + txt, size=9, space_after=3)


# ============================================================
# SECTIONS
# ============================================================
def section_identite(doc, ctx):
    d = ctx.dossier
    superficie = d.get("superficie") or ctx.rapport.get("surface_indicative")
    add_kv_table(doc, [
        ("Par", d.get("demandeur")),
        ("Demeurant à", d.get("demandeur_adresse")),
        ("Sur un terrain sis", d.get("terrain")),
        ("Cadastre", _format_cadastre(d, ctx.rapport)),
        ("Demande déposée le", d.get("date_depot")),
        ("Superficie", f"{superficie} m²" if superficie else None),
        ("N° de dossier", d.get("numero_cu")),
    ])


def section_vu(doc, ctx):
    c = ctx.config
    add_title_bar(doc, "Demande en vue de connaître les dispositions d'urbanisme applicables au terrain")
    for ligne in [
        "Vu la demande de certificat d'urbanisme ;",
        "Vu le code de l'urbanisme ;",
        f"Vu le Plan local d'urbanisme {c.plu_mention} ;",
        f"Vu le Plan de Prévention des Risques Naturels Prévisibles {c.pprn_mention} ;",
        f"Vu le {c.pgri_mention}.",
    ]:
        add_para(doc, ligne)
    doc.add_paragraph()


def section_dpu(doc, ctx):
    add_title_bar(doc, "Droit de préemption")
    # Le DPU vient de infos_surf (ligne libelle='Droit de Préemption Urbain', taguée)
    dpu_txt = None
    for key in ("infos_surf",):
        layer = ctx.rapport.get("intersections", {}).get(key, {})
        for obj in layer.get("objets", []):
            lib = (obj.get("libelle") or "").lower()
            if "préemption" in lib or "preemption" in lib:
                dpu_txt = texte_objet(obj)
                break
    if dpu_txt:
        add_para(doc, dpu_txt)
    else:
        add_para(doc, "La parcelle n'est pas soumise au Droit de Préemption Urbain (DPU).")
    doc.add_paragraph()


def section_sup(doc, ctx):
    groupes = couches_par_section(ctx.rapport)
    layers = groupes.get("sup", [])
    if not layers:
        return
    add_title_bar(doc, "Servitudes d'utilité publique")
    for key, layer in layers:
        _bloc_couche(doc, layer)
    doc.add_paragraph()


def section_risques(doc, ctx):
    groupes = couches_par_section(ctx.rapport)
    layers = groupes.get("risques", [])
    if not layers:
        return
    add_title_bar(doc, "Risques naturels et technologiques")
    for key, layer in layers:
        _bloc_couche(doc, layer)
    doc.add_paragraph()


def section_prescriptions(doc, ctx):
    groupes = couches_par_section(ctx.rapport)
    layers = groupes.get("prescriptions", [])
    if not layers:
        return
    add_title_bar(doc, "Prescriptions et informations applicables au terrain")
    for key, layer in layers:
        _bloc_couche(doc, layer, objets=_objets_affichables(key, layer))
    doc.add_paragraph()


def section_metier(doc, ctx):
    """Modules métier (ENEDIS, Natura/Prairies) : rendu via leurs champs dédiés."""
    inter = ctx.rapport.get("intersections", {})

    enedis = inter.get("reseaux_enedis_lineaires", {})
    analyses = enedis.get("analyses") or []
    if analyses:
        add_title_bar(doc, "Desserte par les réseaux (ENEDIS)")
        for a in analyses:
            diag = a.get("diagnostic_expert_raccordement")
            if diag:
                add_para(doc, f"• {a.get('type_reseau', 'Réseau')} : {diag}", size=9, space_after=3)
        doc.add_paragraph()

    pn = inter.get("prairies_et_natura_2000", {})
    if pn.get("has_natura") or pn.get("has_prairie"):
        add_title_bar(doc, "Natura 2000 / Prairies sensibles")
        diag = pn.get("diagnostic_metier")
        if diag:
            add_para(doc, diag, size=9)
        doc.add_paragraph()


def section_dispositions(doc, ctx):
    c = ctx.config
    add_title_bar(doc, "Nature des dispositions d'urbanisme applicables au terrain")
    for ligne in [
        "Articles d'ordre public du Règlement National d'Urbanisme : R.111-2, R.111-4, "
        "R.111-25 à R.111-27 et R.111-31 à R.111-51 du Code de l'urbanisme.",
        "Articles L.111-6 à L.111-10 du Code de l'urbanisme.",
        "Loi Littoral n° 86-2 du 3 janvier 1986.",
        "Loi Montagne n° 85-30 du 9 janvier 1985.",
        f"Plan local d'urbanisme (PLU) {c.plu_mention}.",
    ]:
        add_para(doc, ligne)

    # Zone PLU depuis le rapport
    zonage = ctx.rapport.get("intersections", {}).get("zonage_plu", {})
    zones, seen = [], set()
    for obj in zonage.get("objets", []):
        z = obj.get("libelle") or obj.get("zonage_reglement")
        if z and z not in seen:
            seen.add(z); zones.append(z)
    if zones:
        add_para(doc, f"La parcelle est située dans la zone {', '.join(zones)} du PLU.", bold=True)
    else:
        add_para(doc, "Zonage PLU non déterminé pour ce terrain.", italic=True)

    # Couches art. 3 : hauteurs, compléments zonage (libellé long / réglementation)
    groupes = couches_par_section(ctx.rapport)
    for key, layer in groupes.get("dispositions", []):
        if key == "zonage_plu":
            for obj in layer.get("objets", []):
                libelong = (obj.get("libelong") or "").strip()
                if libelong and libelong not in zones:
                    add_para(doc, f"• {libelong}", size=9, space_after=3)
                txt = texte_objet(obj)
                zone_code = (obj.get("libelle") or obj.get("zonage_reglement") or "").strip()
                if txt and txt not in zones and txt != libelong and txt != zone_code:
                    add_para(doc, f"• {txt}", size=9, space_after=3)
            continue
        _bloc_couche(doc, layer, objets=_objets_affichables(key, layer))

    add_para(doc, f"Règlement consultable sur : {c.geoportail_url}")
    doc.add_paragraph()


def section_sursis(doc, ctx):
    add_title_bar(doc, "Sursis à statuer")
    add_para(doc, "La commune peut décider de surseoir à statuer, dans les conditions et délais "
                  "prévus à l'article L.424-1 du code de l'urbanisme, sur les demandes d'autorisation "
                  "qui seraient de nature à compromettre ou rendre plus onéreuse l'exécution du futur PLU.")
    doc.add_paragraph()


def section_taxes(doc, ctx):
    c = ctx.config
    add_title_bar(doc, "Taxes et contributions")
    add_para(doc, "Les taxes et contributions ne peuvent être déterminées précisément qu'à l'examen "
                  "de la demande d'autorisation.")
    add_para(doc, "Fiscalité applicable au terrain :", bold=True)
    add_para(doc, f"– Taxe d'aménagement – part communale (taux : {c.taxe_communale})")
    add_para(doc, f"– Taxe d'aménagement – part départementale (taux : {c.taxe_departementale})")
    add_para(doc, f"– Redevance d'archéologie préventive (taux : {c.rap})")
    add_para(doc, "Participations applicables au terrain :", bold=True)
    add_para(doc, "– Projet Urbain Partenarial")
    add_para(doc, "– Participation pour équipements publics exceptionnels")
    doc.add_paragraph()


def section_formalites(doc, ctx):
    add_title_bar(doc, "Formalités administratives préalables à l'opération")
    add_para(doc, "Préalablement à l'édification de construction ou à la réalisation de l'opération "
                  "projetée, les formalités ci-après devront être accomplies : Permis de Construire, "
                  "Permis d'Aménager, Déclaration Préalable, Permis de Démolir.")
    add_para(doc, "Attention : le non-respect de ces formalités ou l'utilisation du sol en "
                  "méconnaissance des règles indiquées est passible de l'amende prévue à "
                  "l'article L.480-4 du Code de l'urbanisme.", italic=True)
    doc.add_paragraph()


def section_signature(doc, ctx):
    c = ctx.config
    add_para(doc, f"{c.nom.title()}, le {datetime.now().strftime('%d/%m/%Y')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, f"Le Maire, {c.maire}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Pour le Maire, par délégation", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, c.delegation, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    add_para(doc, "Le présent certificat est transmis au représentant de l'État (article L.2131-1 du CGCT).",
             size=9, italic=True)


def section_mentions(doc, ctx):
    for m in ctx.config.mentions_communales:
        add_para(doc, f"– {m}", size=9)
    doc.add_paragraph()


def section_informations(doc, ctx):
    add_title_bar(doc, "Informations")
    add_para(doc, "Durée de validité :", bold=True, size=9)
    add_para(doc, "Le certificat est valable dix-huit mois. Une demande déposée dans ce délai bénéficie "
                  "de la cristallisation des règles d'urbanisme, taxes et participations en vigueur à la "
                  "date du certificat, sauf dispositions de sécurité ou de salubrité publique.", size=9)
    add_para(doc, "Prolongation (article R.410-17 du code de l'urbanisme) :", bold=True, size=9)
    add_para(doc, "Prorogeable par période d'un an, sur demande deux mois avant l'expiration, si les règles "
                  "applicables n'ont pas évolué.", size=9)
    add_para(doc, "Délais et voies de recours :", bold=True, size=9)
    add_para(doc, "Recours contentieux possible devant le tribunal administratif dans les deux mois suivant "
                  "la notification (www.telerecours.fr).", size=9)


# Ordre du document
SECTIONS = [
    section_identite,
    section_vu,
    section_dpu,
    section_sup,
    section_risques,
    section_prescriptions,
    section_metier,
    section_dispositions,
    section_sursis,
    section_taxes,
    section_formalites,
    section_signature,
    section_mentions,
    section_informations,
]


# ============================================================
# POINT D'ENTRÉE
# ============================================================
def _setup_document() -> Document:
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10)
    s = doc.sections[0]
    s.page_height = Cm(29.7); s.page_width = Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Cm(2.0))
    return doc


def build_cua(dossier: dict, rapport: dict, output_path: str,
              config: Optional[CommuneConfig] = None) -> str:
    config = config or CommuneConfig()
    ctx = Context(dossier=dossier, rapport=rapport, config=config)
    doc = _setup_document()
    add_para(doc, "CERTIFICAT D'URBANISME", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "délivré par le Maire au nom de la commune", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for section in SECTIONS:
        section(doc, ctx)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import argparse, json
    ap = argparse.ArgumentParser(description="Builder CUA Argelès (v2 — consomme reglementation)")
    ap.add_argument("--rapport", required=True)
    ap.add_argument("--dossier", default=None)
    ap.add_argument("--output", default="CUA_argeles.docx")
    args = ap.parse_args()
    rapport = json.loads(open(args.rapport, encoding="utf-8").read())
    dossier = json.loads(open(args.dossier, encoding="utf-8").read()) if args.dossier else {}
    print("✅ CUA généré :", build_cua(dossier, rapport, args.output))